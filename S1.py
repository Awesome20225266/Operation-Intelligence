from __future__ import annotations

"""
S1 Portal - Work Order Management & Permit To Work (PTW) System

This module implements:
- View Work Order: Filter and display work orders from Supabase
- Request PTW: Full Electrical PTW form mirroring the Word template
- View Applied PTW: Browse and download previously submitted PTWs

All PTW form fields map 1:1 to Electrical_PTW_TEMPLATE.docx placeholders.
Final output is PDF (converted from filled DOCX template).
"""

import os
from io import BytesIO
import subprocess
import tempfile
import time as _time
from datetime import date, datetime, time, timedelta
from zoneinfo import ZoneInfo
from io import BytesIO
from pathlib import Path
import re
from typing import Any

import pandas as pd
import streamlit as st

# Optional PDF overlay support for S3-approved stamp on downloads (S1/S2/S3 consistency)
try:
    from PyPDF2 import PdfReader, PdfWriter  # type: ignore
    from reportlab.lib.colors import Color
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    _HAS_S1_PDF_STAMP_LIBS = True
except Exception:  # pragma: no cover
    _HAS_S1_PDF_STAMP_LIBS = False
import streamlit.components.v1 as components
from docx import Document

from supabase_link import get_supabase_client


# =============================================================================
# CONSTANTS & MAPPINGS
# =============================================================================

TABLE_WORK_ORDERS = "work_orders"
TABLE_PTW_REQUESTS = "ptw_requests"
TABLE_PTW_TEMPLATES = "ptw_templates"

# Internal (derived) statuses — do not change derive_ptw_status outputs
INTERNAL_STATUSES = ["OPEN", "WIP", "APPROVED", "CLOSED", "REJECTED"]

# Canonical UI labels (display only)
UI_STATUSES = ["Closed", "Approved", "Awaiting Approval", "Open", "Rejected"]

# DB <-> UI mapping (keeps DB schema intact while matching requested UI labels)
DB_STATUS_TO_UI = {
    "OPEN": "Open",
    "WIP": "Awaiting Approval",
    "APPROVED": "Approved",
    "CLOSED": "Closed",
    "REJECTED": "Rejected",
}
UI_STATUS_TO_DB = {
    "Open": ["OPEN"],
    "Awaiting Approval": ["WIP"],
    "Approved": ["APPROVED"],
    "Closed": ["CLOSED"],
    "Rejected": ["REJECTED"],
}

# Default status sort order (derived status; single source of truth)
# Required order: Closed, Approved, Awaiting Approval (WIP), Open, Rejected
STATUS_ORDER = {
    "CLOSED": 0,
    "APPROVED": 1,
    "WIP": 2,
    "OPEN": 3,
    "REJECTED": 4,
}

# Placeholder regex for template processing
_PLACEHOLDER_RE = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")


# =============================================================================
# UX HELPERS (UI-only)
# =============================================================================


def _smooth_progress(prog: Any, start: int, end: int, *, text: str, step_delay_s: float = 0.008) -> None:
    """
    UI-only helper to animate progress smoothly.
    Does NOT change business logic; only makes loading feel responsive.
    """
    start_i = int(start)
    end_i = int(end)
    if end_i < start_i:
        start_i, end_i = end_i, start_i
    for p in range(start_i, end_i + 1):
        prog.progress(p, text=text)
        _time.sleep(step_delay_s)


# =============================================================================
# STATUS DERIVATION (SINGLE SOURCE OF TRUTH)
# =============================================================================


def derive_ptw_status(row: dict) -> str:
    """
    Derive PTW status from date columns (single source of truth).
    
    Priority Rules:
    1. REJECTED - If ANY rejection date exists (highest priority)
    2. CLOSED - If date_s1_closed exists
    3. APPROVED - If s1_created, s2_forwarded, s3_approved all exist (no rejections)
    4. WIP - If s1_created exists (no rejections)
    5. OPEN - Fallback when all PTW date columns are NULL
    
    Args:
        row: Dict with date columns (date_s1_created, date_s2_forwarded, 
             date_s3_approved, date_s2_rejected, date_s3_rejected)
    
    Returns:
        Status string: "REJECTED", "APPROVED", "WIP", or "OPEN"
    """
    # Extract date values (handle both dict and pandas Series)
    s1_created = row.get("date_s1_created")
    s2_forwarded = row.get("date_s2_forwarded")
    s3_approved = row.get("date_s3_approved")
    s2_rejected = row.get("date_s2_rejected")
    s3_rejected = row.get("date_s3_rejected")
    s1_closed = row.get("date_s1_closed")
    
    # Helper to check if value is "truthy" (not None, not NaN, not empty)
    def has_value(val) -> bool:
        if val is None:
            return False
        if isinstance(val, float) and pd.isna(val):
            return False
        if isinstance(val, str) and val.strip() == "":
            return False
        return True
    
    # Priority 1: REJECTED (any rejection date exists)
    if has_value(s2_rejected) or has_value(s3_rejected):
        return "REJECTED"

    # Priority 2: CLOSED
    if has_value(s1_closed):
        return "CLOSED"
    
    # Priority 2: APPROVED (full approval chain complete, no rejections)
    if has_value(s1_created) and has_value(s2_forwarded) and has_value(s3_approved):
        return "APPROVED"
    
    # Priority 3: WIP (PTW started but not fully approved)
    # Case 1: Only s1_created exists
    # Case 2: s1_created and s2_forwarded exist but not s3_approved
    if has_value(s1_created):
        return "WIP"
    
    # Priority 5: OPEN (fallback - no PTW activity)
    return "OPEN"


def _update_work_order_s1_created(work_order_id: str, s1_timestamp: str) -> None:
    """
    Update work_orders.date_s1_created when PTW is submitted.
    Only updates if currently NULL (atomic DB-side check).
    Time must already be IST when passed.
    """
    sb = get_supabase_client(prefer_service_role=True)

    resp = (
        sb.table(TABLE_WORK_ORDERS)
        .update({"date_s1_created": s1_timestamp})
        .eq("work_order_id", work_order_id)
        .is_("date_s1_created", None)  # atomic NULL check
        .execute()
    )

    err = getattr(resp, "error", None)
    if err:
        import logging
        logging.warning(f"Failed to update date_s1_created: {err}")


# =============================================================================
# WORK ORDER FUNCTIONS
# =============================================================================


@st.cache_data(show_spinner=False, ttl=60)
def _list_sites_from_work_orders() -> list[str]:
    """Fetch unique site_name values from Supabase work_orders."""
    sb = get_supabase_client(prefer_service_role=True)
    sites: set[str] = set()
    page_size = 1000
    start = 0

    while True:
        resp = (
            sb.table(TABLE_WORK_ORDERS)
            .select("site_name")
            .order("site_name", desc=False)
            .range(start, start + page_size - 1)
            .execute()
        )
        err = getattr(resp, "error", None)
        if err:
            raise RuntimeError(
                f"Failed to fetch site_name from {TABLE_WORK_ORDERS}. "
                f"This is usually a permissions/RLS issue. Error: {err}"
            )
        rows = getattr(resp, "data", None) or []
        if not rows:
            break
        for r in rows:
            v = r.get("site_name")
            if v is None:
                continue
            s = str(v).strip()
            if s:
                sites.add(s)
        if len(rows) < page_size:
            break
        start += page_size

    return sorted(sites)


@st.cache_data(show_spinner=False, ttl=30)
def _list_work_order_ids_by_date(selected_date: date) -> list[str]:
    """
    Fetch work_order_ids from work_orders where date_planned matches selected_date.
    Only returns OPEN work orders that can have PTW requested.
    """
    sb = get_supabase_client(prefer_service_role=True)
    
    # Query work orders for the selected date (date portion of date_planned)
    date_str = selected_date.isoformat()
    
    resp = (
        sb.table(TABLE_WORK_ORDERS)
        .select("work_order_id,date_s1_created,date_s2_forwarded,date_s3_approved,date_s2_rejected,date_s3_rejected")
        .gte("date_planned", f"{date_str}T00:00:00")
        .lt("date_planned", f"{date_str}T23:59:59")
        .order("work_order_id", desc=False)
        .execute()
    )
    
    err = getattr(resp, "error", None)
    if err:
        raise RuntimeError(f"Failed to fetch work_order_ids: {err}")
    
    rows = getattr(resp, "data", None) or []
    out: list[str] = []
    for r in rows:
        woid = str(r.get("work_order_id") or "").strip()
        if not woid:
            continue
        if derive_ptw_status(r) == "OPEN":
            out.append(woid)
    return out


@st.cache_data(show_spinner=False, ttl=30)
def _build_work_order_display_map(selected_date: date) -> dict[str, str]:
    """
    Returns mapping:
    {
        "WO100 - IS16 - Transformer A": "WO100",
        ...
    }
    Only OPEN (derived) work orders for selected date.
    """
    sb = get_supabase_client(prefer_service_role=True)

    date_str = selected_date.isoformat()

    resp = (
        sb.table(TABLE_WORK_ORDERS)
        .select(
            "work_order_id,location,equipment,"
            "date_s1_created,date_s2_forwarded,date_s3_approved,date_s2_rejected,date_s3_rejected"
        )
        .gte("date_planned", f"{date_str}T00:00:00")
        .lt("date_planned", f"{date_str}T23:59:59")
        .order("work_order_id", desc=False)
        .execute()
    )

    err = getattr(resp, "error", None)
    if err:
        raise RuntimeError(f"Failed to fetch work orders: {err}")

    rows = getattr(resp, "data", None) or []

    display_map: dict[str, str] = {}
    for r in rows:
        if derive_ptw_status(r) != "OPEN":
            continue
        wo = str(r.get("work_order_id", "")).strip()
        loc = str(r.get("location", "")).strip()
        eq = str(r.get("equipment", "")).strip()
        if not wo:
            continue
        label = f"{wo} - {loc} - {eq}".strip(" -")
        display_map[label] = wo

    return display_map


def _get_work_order_details(work_order_id: str) -> dict | None:
    """
    Fetch full details of a work order by ID.
    Returns dict with site_name, location, equipment, date_planned, etc.
    """
    sb = get_supabase_client(prefer_service_role=True)
    
    resp = (
        sb.table(TABLE_WORK_ORDERS)
        .select("*")
        .eq("work_order_id", work_order_id)
        .limit(1)
        .execute()
    )
    
    err = getattr(resp, "error", None)
    if err:
        raise RuntimeError(f"Failed to fetch work order details: {err}")
    
    rows = getattr(resp, "data", None) or []
    if not rows:
        return None
    
    return rows[0]


@st.cache_data(show_spinner=False, ttl=60)
def _list_locations_from_work_orders(
    *,
    site_name: str,
    start_date: date,
    end_date: date,
) -> list[str]:
    """Fetch unique location values filtered by site + date range."""
    sb = get_supabase_client(prefer_service_role=True)
    locations: set[str] = set()
    page_size = 1000
    start = 0

    start_dt = datetime.combine(start_date, datetime.min.time())
    end_dt_exclusive = datetime.combine(end_date + timedelta(days=1), datetime.min.time())

    while True:
        q = (
            sb.table(TABLE_WORK_ORDERS)
            .select("location")
            .eq("site_name", site_name)
            .gte("date_planned", start_dt.isoformat(sep=" ", timespec="seconds"))
            .lt("date_planned", end_dt_exclusive.isoformat(sep=" ", timespec="seconds"))
            .order("location", desc=False)
        )
        resp = q.range(start, start + page_size - 1).execute()

        err = getattr(resp, "error", None)
        if err:
            raise RuntimeError(
                f"Failed to fetch location from {TABLE_WORK_ORDERS}. "
                f"This is usually a permissions/RLS issue. Error: {err}"
            )

        rows = getattr(resp, "data", None) or []
        if not rows:
            break
        for r in rows:
            v = r.get("location")
            if v is None:
                continue
            s = str(v).strip()
            if s:
                locations.add(s)
        if len(rows) < page_size:
            break
        start += page_size

    return sorted(locations)


@st.cache_data(show_spinner=False, ttl=60)
def _list_statuses_from_work_orders(
    *,
    site_name: str,
    start_date: date,
    end_date: date,
) -> list[str]:
    """
    Fetch unique DERIVED status values from work_orders, filtered by site + date range.
    
    Status is derived from date columns (single source of truth).
    """
    sb = get_supabase_client(prefer_service_role=True)
    statuses_internal: set[str] = set()
    page_size = 1000
    offset = 0

    start_dt = datetime.combine(start_date, datetime.min.time())
    end_dt_exclusive = datetime.combine(end_date + timedelta(days=1), datetime.min.time())

    # Fetch date columns for status derivation
    cols = "date_s1_created,date_s2_forwarded,date_s3_approved,date_s2_rejected,date_s3_rejected,date_s1_closed"
    
    while True:
        q = (
            sb.table(TABLE_WORK_ORDERS)
            .select(cols)
            .eq("site_name", site_name)
            .gte("date_planned", start_dt.isoformat(sep=" ", timespec="seconds"))
            .lt("date_planned", end_dt_exclusive.isoformat(sep=" ", timespec="seconds"))
        )
        resp = q.range(offset, offset + page_size - 1).execute()

        err = getattr(resp, "error", None)
        if err:
            raise RuntimeError(
                f"Failed to fetch status from {TABLE_WORK_ORDERS}. "
                f"This is usually a permissions/RLS issue. Error: {err}"
            )

        rows = getattr(resp, "data", None) or []
        if not rows:
            break
        
        for r in rows:
            # Derive status from date columns
            derived = derive_ptw_status(r)
            if derived in INTERNAL_STATUSES:
                statuses_internal.add(derived)
        
        if len(rows) < page_size:
            break
        offset += page_size

    ordered_internal = sorted(statuses_internal, key=lambda s: STATUS_ORDER.get(s, 99))
    # Return UI labels (display only)
    return [DB_STATUS_TO_UI.get(str(s).strip().upper(), str(s)) for s in ordered_internal]


@st.cache_data(show_spinner=False, ttl=60)
def _fetch_work_orders(
    *,
    site_name: str,
    start_date: date,
    end_date: date,
    status_ui: str | None,
    location: str | None,
) -> pd.DataFrame:
    """
    Fetch work orders filtered by site, date range, status, and location.
    
    Status is DERIVED from date columns (not from stored status field).
    """
    sb = get_supabase_client(prefer_service_role=True)

    start_dt = datetime.combine(start_date, datetime.min.time())
    end_dt_exclusive = datetime.combine(end_date + timedelta(days=1), datetime.min.time())

    # Include date columns for status derivation
    cols = (
        "work_order_id,location,equipment,frequency,isolation_requirement,date_planned,"
        "date_s1_created,date_s2_forwarded,date_s3_approved,date_s2_rejected,date_s3_rejected,date_s1_closed"
    )
    q = sb.table(TABLE_WORK_ORDERS).select(cols).eq("site_name", site_name)

    q = q.gte("date_planned", start_dt.isoformat(sep=" ", timespec="seconds")).lt(
        "date_planned", end_dt_exclusive.isoformat(sep=" ", timespec="seconds")
    )

    if location:
        q = q.eq("location", location)
    
    # NOTE: We no longer filter by stored status - we derive status from dates
    # and filter AFTER deriving

    q = q.order("date_planned", desc=False)

    resp = q.execute()
    err = getattr(resp, "error", None)
    if err:
        raise RuntimeError(f"Failed to fetch work orders: {err}")

    data: list[dict[str, Any]] = getattr(resp, "data", None) or []
    df = pd.DataFrame(data)
    
    if df.empty:
        # Return empty DataFrame with expected columns
        return pd.DataFrame(columns=[
            "work_order_id", "location", "equipment", "frequency",
            "isolation_requirement", "date_planned", "status"
        ])
    
    # DERIVE status from date columns (single source of truth)
    df["status"] = df.apply(lambda row: derive_ptw_status(row.to_dict()), axis=1)
    
    # Filter by derived status if requested (derived statuses only)
    if status_ui:
        sel = str(status_ui).strip()
        internal_targets = UI_STATUS_TO_DB.get(sel, [sel.strip().upper()])
        target = internal_targets[0] if internal_targets else sel.strip().upper()
        if target == "WIP":
            df = df[df["status"] == "WIP"]
        else:
            df = df[df["status"] == target]
    
    # Select only display columns (hide internal date columns)
    ordered_cols = [
        "work_order_id",
        "location",
        "equipment",
        "frequency",
        "isolation_requirement",
        "date_planned",
        "status",
    ]
    for c in ordered_cols:
        if c not in df.columns:
            df[c] = None
    df = df[ordered_cols]

    if not df.empty:
        df["date_planned"] = pd.to_datetime(df["date_planned"], errors="coerce").dt.date.astype("string")
        df["_status_rank"] = df["status"].astype("string").fillna("").map(
            lambda s: STATUS_ORDER.get(str(s).strip().upper(), 99)
        )
        df = df.sort_values(by=["_status_rank", "date_planned"], ascending=[True, True]).drop(columns=["_status_rank"])

    return df


def _highlight_status(val: Any) -> str:
    """
    Apply color coding to status cell based on derived status.
    
    Colors:
    - OPEN: Blue (#2563eb)
    - WIP: Orange (#f97316)
    - APPROVED: Green (#10b981)
    - CLOSED: Darker green (#065f46)
    - REJECTED: Red (#dc2626)
    """
    status = str(val).strip().upper()
    
    if status == "OPEN":
        return "background-color: #2563eb; color: white; font-weight: 700;"
    elif status in ("WIP", "AWAITING APPROVAL"):
        return "background-color: #f97316; color: white; font-weight: 700;"
    elif status == "APPROVED":
        return "background-color: #10b981; color: white; font-weight: 700;"
    elif status == "CLOSED":
        return "background-color: #065f46; color: white; font-weight: 700;"
    elif status == "REJECTED":
        return "background-color: #dc2626; color: white; font-weight: 700;"
    return ""


# Keep old name for backward compatibility
_highlight_open_status = _highlight_status


# =============================================================================
# PTW FUNCTIONS
# =============================================================================

# Template storage bucket name in Supabase
SUPABASE_TEMPLATE_BUCKET = "ptw-templates"
TEMPLATE_FILE_NAME = "Electrical_PTW_TEMPLATE.docx"
TEMPLATE_PDF_FILE_NAME = "PDF_Electrical_PTW_TEMPLATE.pdf"  # ReportLab overlay template


def _ptw_exists_for_work_orders(selected_ids: list[str]) -> tuple[bool, str | None]:
    """
    Prevent duplicate PTW coverage.
    If ANY selected work_order_id already exists in any PTW, block submission.

    Backward compatible with legacy single-WO PTWs.
    """
    if not selected_ids:
        return False, None

    sb = get_supabase_client(prefer_service_role=True)

    resp = sb.table(TABLE_PTW_REQUESTS).select("ptw_id,permit_no,form_data").order("created_at", desc=True).execute()
    err = getattr(resp, "error", None)
    if err:
        raise RuntimeError(f"Failed to check existing PTWs: {err}")

    rows = getattr(resp, "data", None) or []
    selected_set = {str(x).strip() for x in selected_ids if str(x).strip()}
    if not selected_set:
        return False, None

    for row in rows:
        if not isinstance(row, dict):
            continue
        fd = row.get("form_data") or {}
        if not isinstance(fd, dict):
            fd = {}

        existing_ids = fd.get("work_order_ids")
        if isinstance(existing_ids, list) and existing_ids:
            existing_set = {str(x).strip() for x in existing_ids if str(x).strip()}
            if existing_set & selected_set:
                return True, row.get("ptw_id")

        # Legacy support: single-work-order PTWs
        permit_no = str(row.get("permit_no") or "").strip()
        legacy_wo = str(fd.get("work_order_id") or "").strip()
        if permit_no and permit_no in selected_set:
            return True, row.get("ptw_id")
        if legacy_wo and legacy_wo in selected_set:
            return True, row.get("ptw_id")

    return False, None


def _ptw_exists_for_work_order(work_order_id: str) -> tuple[bool, str | None]:
    """Backward compatible single-WO wrapper."""
    return _ptw_exists_for_work_orders([work_order_id] if work_order_id else [])


def _reset_ptw_form_state() -> None:
    """
    Fully reset PTW form inputs after SUCCESSFUL submission.

    Streamlit widget state is keyed; changing the form key alone does NOT reset widget keys.
    We explicitly clear PTW-related widget keys so the UI visually resets without a full-page feel.
    """
    # Prefixes for form input keys to clear
    prefixes = ("hz_", "rk_", "ppe_", "sp_", "ap_", "chk_", "ptw_")
    
    # Keys to explicitly keep (don't delete)
    keep = {
        # Keep download + view state for the success screen
        "s1_ptw_last_file",
        "s1_ptw_last_permit_no",
        "s1_ptw_last_ext",
        "s1_ptw_view_df",
        "s1_ptw_reset_counter",
        "ptw_just_submitted",  # Keep this for success screen
        # Keep S1 portal state
        "nav_page",
    }

    # Clear all form-related keys
    keys_to_delete = []
    for k in st.session_state.keys():
        if k in keep:
            continue
        if k.startswith(prefixes):
            keys_to_delete.append(k)
    
    for k in keys_to_delete:
        try:
            del st.session_state[k]
        except Exception:
            pass
    
    # Increment reset counter to force new form key
    st.session_state["s1_ptw_reset_counter"] = st.session_state.get("s1_ptw_reset_counter", 0) + 1


@st.cache_data(show_spinner=False, ttl=3600)
def _download_template_from_supabase() -> bytes:
    """
    Download the Electrical PTW template from Supabase storage bucket.
    
    Returns the template file as bytes.
    Caches for 1 hour to avoid repeated downloads.
    """
    sb = get_supabase_client(prefer_service_role=True)
    
    preferred = (_get_setting("PTW_TEMPLATE_FORMAT") or "auto").strip().lower()
    if preferred in {"docx", "docx_first", "docx-first"}:
        candidates = [TEMPLATE_FILE_NAME, TEMPLATE_PDF_FILE_NAME]
    elif preferred in {"pdf", "pdf_first", "pdf-first"}:
        candidates = [TEMPLATE_PDF_FILE_NAME, TEMPLATE_FILE_NAME]
    else:
        # auto: try PDF first then DOCX
        candidates = [TEMPLATE_PDF_FILE_NAME, TEMPLATE_FILE_NAME]

    last_err: Exception | None = None
    for fname in candidates:
        try:
            response = sb.storage.from_(SUPABASE_TEMPLATE_BUCKET).download(fname)
            if isinstance(response, (bytes, bytearray)) and len(response) > 0:
                return bytes(response)
            raise RuntimeError("Empty template file received from Supabase storage")
        except Exception as e:
            last_err = e
            continue

    raise RuntimeError(
        f"Failed to download template from Supabase storage.\n"
        f"Bucket: {SUPABASE_TEMPLATE_BUCKET}\n"
        f"Tried: {', '.join(candidates)}\n"
        f"Error: {last_err}\n\n"
        f"Please ensure:\n"
        f"1. The bucket '{SUPABASE_TEMPLATE_BUCKET}' exists in Supabase Storage\n"
        f"2. At least one of these files is uploaded: {', '.join(candidates)}\n"
        f"3. Your service role key has read access to the bucket"
    ) from last_err


def _tick(val: bool | str | None) -> str:
    """
    Convert boolean/truthy value to BOLD checkmark for document.
    
    Industrial PTW standard:
    - Selected checkbox -> BOLD tick mark (checkmark symbol)
    - Unselected checkbox -> blank (empty string)
    """
    if val is None:
        return ""
    if isinstance(val, bool):
        return "\u2713" if val else ""  # Unicode checkmark
    if isinstance(val, str):
        return "\u2713" if val.upper() in ("Y", "YES", "TRUE", "1") else ""
    return "\u2713" if bool(val) else ""


def insert_ptw_request(
    *,
    permit_no: str,
    site_name: str,
    created_by: str,
    form_data: dict,
) -> str:
    """Insert PTW request into Supabase and return the ptw_id."""
    sb = get_supabase_client(prefer_service_role=True)

    # Fetch active electrical template
    tpl = (
        sb.table(TABLE_PTW_TEMPLATES)
        .select("template_id")
        .eq("permit_type", "ELECTRICAL")
        .eq("is_active", True)
        .limit(1)
        .execute()
    )

    tpl_err = getattr(tpl, "error", None)
    if tpl_err:
        raise RuntimeError(tpl_err)

    tpl_data = getattr(tpl, "data", None) or []
    if not tpl_data:
        raise RuntimeError("No active Electrical PTW template found in database")

    template_id = tpl_data[0]["template_id"]

    resp = (
        sb.table(TABLE_PTW_REQUESTS)
        .insert(
            {
                "template_id": template_id,
                "permit_type": "ELECTRICAL",
                "permit_no": permit_no,
                "site_name": site_name,
                "status": "SUBMITTED",
                "form_data": form_data,
                "created_by": created_by,
            }
        )
        .execute()
    )

    resp_err = getattr(resp, "error", None)
    if resp_err:
        raise RuntimeError(resp_err)

    resp_data = getattr(resp, "data", None) or []
    if not resp_data:
        raise RuntimeError("PTW insert succeeded but returned no data")

    return resp_data[0]["ptw_id"]


@st.cache_data(show_spinner=False, ttl=60)
def fetch_ptw_requests(
    *,
    start_date: date | None = None,
    end_date: date | None = None,
) -> pd.DataFrame:
    """
    Fetch PTW requests from Supabase, optionally filtered by date range.
    
    Status is DERIVED from work_orders date columns (single source of truth).
    """
    sb = get_supabase_client(prefer_service_role=True)

    # IMPORTANT: Date filtering for lifecycle visibility is based on work_orders.date_s1_created,
    # NOT ptw_requests.created_at (ptw_requests.created_at is not the lifecycle clock).
    q = sb.table(TABLE_PTW_REQUESTS).select("ptw_id,permit_no,site_name,status,created_at,created_by,form_data")
    q = q.order("created_at", desc=True)
    resp = q.execute()

    err = getattr(resp, "error", None)
    if err:
        raise RuntimeError(f"Failed to fetch PTW requests: {err}")

    data = getattr(resp, "data", None) or []
    df = pd.DataFrame(data)
    
    if df.empty:
        return df
    
    # Derive status + lifecycle dates from work_orders date columns.
    # Multi-WO PTWs: ids are stored in form_data["work_order_ids"] (fallback to permit_no legacy).
    def _row_work_order_ids(r: dict) -> list[str]:
        try:
            fd = r.get("form_data") or {}
            if isinstance(fd, dict):
                ids = fd.get("work_order_ids")
                if isinstance(ids, list) and ids:
                    return [str(x).strip() for x in ids if str(x).strip()]
                legacy = str(fd.get("work_order_id") or "").strip()
                if legacy:
                    return [legacy]
        except Exception:
            pass
        pn = str(r.get("permit_no") or "").strip()
        return [pn] if pn else []

    all_ids: set[str] = set()
    ids_per_row: list[list[str]] = []
    for _, row in df.iterrows():
        rr = row.to_dict() if hasattr(row, "to_dict") else {}
        ids = _row_work_order_ids(rr)
        ids_per_row.append(ids)
        all_ids.update(ids)

    status_lookup: dict[str, str] = {}
    s1_created_lookup: dict[str, Any] = {}
    if all_ids:
        wo_resp = (
            sb.table(TABLE_WORK_ORDERS)
            .select(
                "work_order_id,date_s1_created,date_s2_forwarded,date_s3_approved,"
                "date_s2_rejected,date_s3_rejected,date_s1_closed"
            )
            .in_("work_order_id", list(all_ids))
            .execute()
        )
        wo_data = getattr(wo_resp, "data", None) or []
        if wo_data:
            wo_df = pd.DataFrame(wo_data)
            for _, row in wo_df.iterrows():
                derived = derive_ptw_status(row.to_dict())
                woid = str(row.get("work_order_id") or "").strip()
                if not woid:
                    continue
                status_lookup[woid] = derived
                s1_created_lookup[woid] = row.get("date_s1_created")

    def _aggregate_status(statuses: list[str]) -> str:
        sset = {str(s).strip().upper() for s in statuses if str(s).strip()}
        if "REJECTED" in sset:
            return "REJECTED"
        if sset and sset.issubset({"CLOSED"}):
            return "CLOSED"
        if sset and sset.issubset({"APPROVED"}):
            return "APPROVED"
        if sset and sset.issubset({"OPEN"}):
            return "OPEN"
        # Any mix defaults to WIP (submitted but not fully approved/closed)
        return "WIP"

    derived_statuses: list[str] = []
    derived_s1_dates: list[Any] = []
    for ids in ids_per_row:
        sts = [status_lookup.get(i, "WIP") for i in ids] if ids else ["WIP"]
        derived_statuses.append(_aggregate_status(sts))
        # Use earliest non-null s1_created for date filtering/display
        dates = [s1_created_lookup.get(i) for i in ids if s1_created_lookup.get(i) is not None]
        derived_s1_dates.append(min(dates) if dates else None)

    df["status"] = derived_statuses
    df["date_s1_created"] = derived_s1_dates

    # Apply lifecycle date range filter using work_orders.date_s1_created
    if start_date or end_date:
        start_dt = datetime.combine(start_date or date.min, datetime.min.time())
        end_dt_excl = datetime.combine((end_date or date.max) + timedelta(days=1), datetime.min.time())
        ts = pd.to_datetime(df.get("date_s1_created"), errors="coerce")
        df = df[(ts >= start_dt) & (ts < end_dt_excl)]
        df = df.sort_values(by=["date_s1_created"], ascending=[False], na_position="last")
    
    return df


def generate_ptw_docx_from_template(template_bytes: bytes, form_data: dict) -> bytes:
    """
    Fill Electrical PTW Word template with form data.
    
    Processing rules:
    - Replace {{key}} with value if present in form_data
    - Replace {{key}} with empty string if not present (no leftover placeholders)
    - User-input values appear in GREEN BOLD
    
    Args:
        template_bytes: The template DOCX file as bytes (from Supabase storage)
        form_data: Dictionary of placeholder keys to values
    
    Returns in-memory DOCX file as bytes.
    """
    from docx.shared import RGBColor
    
    mapping: dict[str, str] = {}
    for k, v in (form_data or {}).items():
        mapping[str(k)] = "" if v is None else str(v)

    # Load document from bytes
    doc = Document(BytesIO(template_bytes))
    
    # Green color for user inputs
    GREEN_COLOR = RGBColor(0, 128, 0)
    
    # More aggressive placeholder pattern - handles various bracket styles
    placeholder_pattern = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")

    def get_full_paragraph_text(paragraph) -> str:
        """Get full text by explicitly concatenating all runs."""
        runs = getattr(paragraph, "runs", None)
        if runs:
            return "".join(r.text or "" for r in runs)
        return getattr(paragraph, "text", "") or ""

    def replace_in_paragraph_with_formatting(paragraph) -> None:
        """
        Replace placeholders with GREEN BOLD user values.
        
        Strategy:
        1. Get full text from all runs (handles split placeholders)
        2. Find and replace all placeholders via regex
        3. Completely remove old runs and rebuild paragraph
        """
        # Get full text by concatenating all runs
        full_text = get_full_paragraph_text(paragraph)
        
        # Also try paragraph.text as backup
        if not full_text:
            full_text = getattr(paragraph, "text", "") or ""
        
        # Skip if truly empty
        if not full_text:
            return
        
        # Always try to find placeholders (don't skip based on "{" check)
        matches = list(placeholder_pattern.finditer(full_text))
        if not matches:
            # No placeholders found, nothing to do
            return
        
        # Build parts list: (type, text) where type is "static" or "user"
        parts = []
        last_end = 0
        
        for m in matches:
            # Add text before this placeholder
            if m.start() > last_end:
                parts.append(("static", full_text[last_end:m.start()]))
            
            # Get replacement value
            key = m.group(1).strip()
            replacement = mapping.get(key, "")
            
            if replacement:
                parts.append(("user", replacement))
            # If no replacement, placeholder is removed (empty string)
            
            last_end = m.end()
        
        # Add remaining text after last placeholder
        if last_end < len(full_text):
            parts.append(("static", full_text[last_end:]))
        
        # COMPLETELY clear the paragraph by removing all run elements
        p_element = paragraph._element
        # Find and remove all 'w:r' (run) elements
        for child in list(p_element):
            tag_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag_name == 'r':
                p_element.remove(child)
        
        # Add new runs with appropriate formatting
        for part_type, part_text in parts:
            if not part_text:
                continue
            new_run = paragraph.add_run(part_text)
            if part_type == "user":
                new_run.bold = True
                new_run.font.color.rgb = GREEN_COLOR

    def replace_in_cell(cell) -> None:
        """Process all paragraphs and nested tables in a cell."""
        for p in cell.paragraphs:
            replace_in_paragraph_with_formatting(p)
        for t in cell.tables:
            replace_in_table(t)

    def replace_in_table(table) -> None:
        """Process all cells in a table."""
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell)

    # Process body paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph_with_formatting(p)

    # Process tables (template is table-driven)
    for t in doc.tables:
        replace_in_table(t)

    # Process headers/footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph_with_formatting(p)
        for t in section.header.tables:
            replace_in_table(t)
        for p in section.footer.paragraphs:
            replace_in_paragraph_with_formatting(p)
        for t in section.footer.tables:
            replace_in_table(t)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _get_setting(name: str) -> str | None:
    """Read from Streamlit secrets when available; fallback to environment variables."""
    try:
        import streamlit as _st  # type: ignore

        v = _st.secrets.get(name)  # type: ignore[attr-defined]
        if v is not None and str(v).strip() != "":
            return str(v)
    except Exception:
        pass
    v = os.getenv(name)
    return str(v) if v is not None and str(v).strip() != "" else None


def _convert_docx_to_pdf_cloudconvert(docx_bytes: bytes, *, timeout_s: int = 180, progress_callback=None) -> bytes | None:
    """
    CloudConvert DOCX→PDF conversion (robust for hosted deployments).

    Requires: CLOUDCONVERT_API_KEY in Streamlit secrets or environment variables.
    Returns PDF bytes on success, or None if not configured / failed.
    """
    # Support multiple secret names for convenience in hosted environments
    api_key = (
        _get_setting("CLOUDCONVERT_API_KEY")
        or _get_setting("CLOUD_CONVERT_API_KEY")
        or _get_setting("CLOUDCONVERT_KEY")
    )
    if not api_key:
        return None

    try:
        import time
        import requests

        headers = {"Authorization": f"Bearer {api_key}"}

        if progress_callback:
            try:
                progress_callback(62, "Starting cloud PDF conversion...")
            except Exception:
                pass

        # Create a job: import/upload -> convert -> export/url
        job_req = {
            "tasks": {
                "import-ptw": {"operation": "import/upload"},
                "convert-ptw": {
                    "operation": "convert",
                    "input": "import-ptw",
                    "output_format": "pdf",
                },
                "export-ptw": {"operation": "export/url", "input": "convert-ptw"},
            }
        }

        r = requests.post(
            "https://api.cloudconvert.com/v2/jobs",
            headers={**headers, "Content-Type": "application/json"},
            json=job_req,
            timeout=30,
        )
        r.raise_for_status()
        job = r.json().get("data") or {}
        job_id = job.get("id")
        if not job_id:
            return None

        # Poll until import/upload task contains the upload form (CloudConvert can be async)
        deadline = time.time() + max(30, int(timeout_s))
        upload_url = None
        upload_params: dict | None = None
        while time.time() < deadline and (not upload_url or not upload_params):
            jr0 = requests.get(f"https://api.cloudconvert.com/v2/jobs/{job_id}", headers=headers, timeout=30)
            jr0.raise_for_status()
            job0 = jr0.json().get("data") or {}
            tasks0 = job0.get("tasks") or []
            import_task = next((t for t in tasks0 if t.get("name") == "import-ptw"), None) or {}
            form = (import_task.get("result") or {}).get("form") or {}
            upload_url = form.get("url")
            upload_params = form.get("parameters") or None
            if upload_url and upload_params:
                break
            time.sleep(0.4)

        if not upload_url or not upload_params:
            return None

        # Upload file to CloudConvert
        if progress_callback:
            try:
                progress_callback(70, "Uploading document to converter...")
            except Exception:
                pass
        files = {
            "file": (
                "ptw.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        }
        ur = requests.post(upload_url, data=upload_params, files=files, timeout=60)
        ur.raise_for_status()

        # Poll job status
        export_url: str | None = None
        while time.time() < deadline:
            jr = requests.get(f"https://api.cloudconvert.com/v2/jobs/{job_id}", headers=headers, timeout=30)
            jr.raise_for_status()
            job_data = jr.json().get("data") or {}
            status = job_data.get("status")
            tasks2 = job_data.get("tasks") or []

            if status == "error":
                return None
            if status == "finished":
                export_task = next((t for t in tasks2 if t.get("name") == "export-ptw"), None) or {}
                files_out = ((export_task.get("result") or {}).get("files")) or []
                if files_out and isinstance(files_out, list) and files_out[0].get("url"):
                    export_url = files_out[0].get("url")
                    break

            time.sleep(0.8)

        if not export_url:
            return None

        if progress_callback:
            try:
                progress_callback(90, "Downloading PDF...")
            except Exception:
                pass
        pdf_r = requests.get(export_url, timeout=60)
        pdf_r.raise_for_status()
        pdf_bytes = pdf_r.content
        if isinstance(pdf_bytes, (bytes, bytearray)) and pdf_bytes[:4] == b"%PDF":
            if progress_callback:
                try:
                    progress_callback(98, "PDF ready")
                except Exception:
                    pass
            return bytes(pdf_bytes)
        return None
    except Exception:
        return None


# =============================================================================
# PDF TEMPLATE + REPORTLAB OVERLAY (NO DOCX/LIBREOFFICE/CLOUDCONVERT)
# =============================================================================

# Authoritative PDF overlay coordinate mapping (A4: 595 x 842, bottom-left origin).
# Structure: dotted_tag -> {page, x, y, width, height, type}
# Types: "text" | "multiline" | "checkbox" | "tri_state"
PTW_PDF_COORDINATES: dict[str, dict[str, Any]] = {
    # ---------------------------
    # Page 1 — Header (text)
    # ---------------------------
    "meta.permit_no": {"page": 1, "x": 166, "y": 666, "width": 181, "height": 19, "type": "text"},
    "meta.permit_validity_date": {"page": 1, "x": 461, "y": 665, "width": 118, "height": 20, "type": "text"},
    "meta.start_time": {"page": 1, "x": 166, "y": 644, "width": 184, "height": 16, "type": "text"},
    "meta.end_time": {"page": 1, "x": 458, "y": 646, "width": 123, "height": 17, "type": "text"},
    "meta.project_site_name": {"page": 1, "x": 163, "y": 626, "width": 188, "height": 15, "type": "text"},
    "meta.work_location": {"page": 1, "x": 456, "y": 624, "width": 121, "height": 19, "type": "text"},
    "meta.description_of_work": {"page": 1, "x": 158, "y": 605, "width": 422, "height": 16, "type": "text"},
    "meta.contractor_name": {"page": 1, "x": 161, "y": 582, "width": 419, "height": 17, "type": "text"},
    # ---------------------------
    # Page 1 — Hazards (checkboxes + text)
    # ---------------------------
    "hazards.live_dc_cables": {"page": 1, "x": 34, "y": 529, "width": 9, "height": 11, "type": "checkbox"},
    "hazards.loose_or_damaged_connectors": {"page": 1, "x": 135, "y": 528, "width": 14, "height": 15, "type": "checkbox"},
    "hazards.moving_solar_tracker_parts": {"page": 1, "x": 288, "y": 531, "width": 11, "height": 10, "type": "checkbox"},
    "hazards.dust_accumulation_on_equipment": {"page": 1, "x": 417, "y": 529, "width": 15, "height": 13, "type": "checkbox"},
    "hazards.high_dc_voltage_from_pv_panels": {"page": 1, "x": 33, "y": 512, "width": 15, "height": 14, "type": "checkbox"},
    "hazards.poor_grounding_earthing": {"page": 1, "x": 135, "y": 507, "width": 14, "height": 14, "type": "checkbox"},
    "hazards.heavy_solar_panels": {"page": 1, "x": 288, "y": 507, "width": 11, "height": 12, "type": "checkbox"},
    "hazards.wildlife_snakes_insects_birds": {"page": 1, "x": 419, "y": 506, "width": 13, "height": 14, "type": "checkbox"},
    "hazards.arc_flash_short_circuit": {"page": 1, "x": 34, "y": 489, "width": 15, "height": 13, "type": "checkbox"},
    "hazards.working_at_height_rooftop_structure": {"page": 1, "x": 135, "y": 487, "width": 16, "height": 14, "type": "checkbox"},
    "hazards.sharp_panel_edges_metal_structures": {"page": 1, "x": 287, "y": 487, "width": 14, "height": 12, "type": "checkbox"},
    "hazards.lightning_or_stormy_weather": {"page": 1, "x": 417, "y": 479, "width": 15, "height": 17, "type": "checkbox"},
    "hazards.improper_grounding": {"page": 1, "x": 34, "y": 462, "width": 15, "height": 13, "type": "checkbox"},
    "hazards.wet_surfaces": {"page": 1, "x": 135, "y": 453, "width": 17, "height": 20, "type": "checkbox"},
    "hazards.high_ambient_temperature_sun_exposure": {"page": 1, "x": 285, "y": 463, "width": 15, "height": 12, "type": "checkbox"},
    "hazards.overloaded_circuits": {"page": 1, "x": 417, "y": 464, "width": 14, "height": 15, "type": "checkbox"},
    "hazards.manual_handling": {"page": 1, "x": 136, "y": 436, "width": 13, "height": 12, "type": "checkbox"},
    "hazards.overhead_line": {"page": 1, "x": 286, "y": 436, "width": 14, "height": 10, "type": "checkbox"},
    "hazards.if_others_text": {"page": 1, "x": 537, "y": 434, "width": 45, "height": 15, "type": "text"},
    # ---------------------------
    # Page 1 — Risks (checkboxes + multiline text)
    # ---------------------------
    "risks.electrocution": {"page": 1, "x": 31, "y": 388, "width": 15, "height": 15, "type": "checkbox"},
    "risks.burns_equipment_damage": {"page": 1, "x": 135, "y": 387, "width": 14, "height": 15, "type": "checkbox"},
    "risks.unexpected_energization": {"page": 1, "x": 286, "y": 387, "width": 14, "height": 15, "type": "checkbox"},
    "risks.heat_stress_dehydration": {"page": 1, "x": 286, "y": 387, "width": 14, "height": 15, "type": "checkbox"},
    "risks.electric_shock": {"page": 1, "x": 35, "y": 367, "width": 11, "height": 14, "type": "checkbox"},
    "risks.fire_or_overheating": {"page": 1, "x": 135, "y": 367, "width": 15, "height": 12, "type": "checkbox"},
    "risks.crushing_or_pinching_injuries": {"page": 1, "x": 285, "y": 370, "width": 18, "height": 17, "type": "checkbox"},
    "risks.if_others_text": {"page": 1, "x": 430, "y": 330, "width": 149, "height": 43, "type": "multiline"},
    "risks.electric_burn": {"page": 1, "x": 31, "y": 345, "width": 18, "height": 17, "type": "checkbox"},
    "risks.fall_causing_serious_injury": {"page": 1, "x": 135, "y": 346, "width": 14, "height": 13, "type": "checkbox"},
    "risks.back_or_muscle_injury": {"page": 1, "x": 288, "y": 348, "width": 12, "height": 11, "type": "checkbox"},
    "risks.bites_stings": {"page": 1, "x": 29, "y": 327, "width": 17, "height": 15, "type": "checkbox"},
    "risks.falling_particles": {"page": 1, "x": 136, "y": 327, "width": 14, "height": 15, "type": "checkbox"},
    "risks.tripping_slipping": {"page": 1, "x": 289, "y": 329, "width": 10, "height": 11, "type": "checkbox"},
    # ---------------------------
    # Page 1 — PPE (checkboxes + text)
    # ---------------------------
    "ppe.safety_helmet": {"page": 1, "x": 31, "y": 271, "width": 15, "height": 15, "type": "checkbox"},
    "ppe.hrc_suit": {"page": 1, "x": 135, "y": 272, "width": 15, "height": 14, "type": "checkbox"},
    "ppe.respirators": {"page": 1, "x": 286, "y": 272, "width": 14, "height": 11, "type": "checkbox"},
    "ppe.full_body_safety_harness": {"page": 1, "x": 417, "y": 272, "width": 14, "height": 15, "type": "checkbox"},
    "ppe.safety_shoes": {"page": 1, "x": 32, "y": 254, "width": 13, "height": 13, "type": "checkbox"},
    "ppe.electrical_mat": {"page": 1, "x": 138, "y": 255, "width": 14, "height": 14, "type": "checkbox"},
    "ppe.dust_masks": {"page": 1, "x": 287, "y": 253, "width": 15, "height": 16, "type": "checkbox"},
    "ppe.lifeline_anchor_point": {"page": 1, "x": 417, "y": 253, "width": 15, "height": 15, "type": "checkbox"},
    "ppe.reflective_vest": {"page": 1, "x": 30, "y": 236, "width": 17, "height": 14, "type": "checkbox"},
    "ppe.face_shield_with_arc_protection": {"page": 1, "x": 135, "y": 237, "width": 14, "height": 15, "type": "checkbox"},
    "ppe.ear_plugs_muffs": {"page": 1, "x": 287, "y": 237, "width": 15, "height": 15, "type": "checkbox"},
    "ppe.cut_resistant_gloves": {"page": 1, "x": 417, "y": 237, "width": 13, "height": 14, "type": "checkbox"},
    "ppe.safety_goggles": {"page": 1, "x": 31, "y": 218, "width": 16, "height": 14, "type": "checkbox"},
    "ppe.insulated_hand_tools": {"page": 1, "x": 135, "y": 219, "width": 14, "height": 17, "type": "checkbox"},
    "ppe.electrical_hand_gloves": {"page": 1, "x": 286, "y": 216, "width": 14, "height": 17, "type": "checkbox"},
    "ppe.if_others_text": {"page": 1, "x": 418, "y": 221, "width": 13, "height": 15, "type": "text"},
    # ---------------------------
    # Page 1 — Safety precautions (checkboxes + multiline)
    # ---------------------------
    "precautions.electrical_isolation": {"page": 1, "x": 32, "y": 160, "width": 15, "height": 19, "type": "checkbox"},
    "precautions.fire_extinguisher_first_aid": {"page": 1, "x": 134, "y": 158, "width": 18, "height": 16, "type": "checkbox"},
    "precautions.proper_isolation": {"page": 1, "x": 285, "y": 159, "width": 17, "height": 21, "type": "checkbox"},
    "precautions.authorized_competent_personnel": {"page": 1, "x": 415, "y": 165, "width": 15, "height": 16, "type": "checkbox"},
    "precautions.lockout_tagout": {"page": 1, "x": 33, "y": 133, "width": 13, "height": 19, "type": "checkbox"},
    "precautions.warning_signage_barricading": {"page": 1, "x": 136, "y": 144, "width": 11, "height": 10, "type": "checkbox"},
    "precautions.rescue_equipment_at_site": {"page": 1, "x": 288, "y": 135, "width": 12, "height": 17, "type": "checkbox"},
    "precautions.if_others_text": {"page": 1, "x": 420, "y": 103, "width": 162, "height": 36, "type": "multiline"},
    "precautions.zero_voltage_ensure": {"page": 1, "x": 35, "y": 115, "width": 11, "height": 15, "type": "checkbox"},
    "precautions.proper_earthing_grounding": {"page": 1, "x": 135, "y": 114, "width": 16, "height": 16, "type": "checkbox"},
    "precautions.pre_job_meeting": {"page": 1, "x": 286, "y": 114, "width": 16, "height": 17, "type": "checkbox"},
    "precautions.insulated_tools": {"page": 1, "x": 33, "y": 102, "width": 13, "height": 12, "type": "checkbox"},
    "precautions.proper_illumination": {"page": 1, "x": 135, "y": 102, "width": 15, "height": 14, "type": "checkbox"},
    "precautions.escape_route_clear": {"page": 1, "x": 283, "y": 100, "width": 19, "height": 14, "type": "checkbox"},
    # ---------------------------
    # Page 2 — Associated permits / documents (checkboxes + text)
    # ---------------------------
    "permits.hot_work_permit_no": {"page": 2, "x": 31, "y": 640, "width": 15, "height": 14, "type": "checkbox"},
    "permits.work_at_night_permit_no": {"page": 2, "x": 151, "y": 639, "width": 17, "height": 18, "type": "checkbox"},
    "permits.work_at_height_permit_no": {"page": 2, "x": 328, "y": 638, "width": 17, "height": 16, "type": "checkbox"},
    "permits.general_work_permit_no": {"page": 2, "x": 29, "y": 621, "width": 22, "height": 16, "type": "checkbox"},
    "permits.excavation_work_permit_no": {"page": 2, "x": 152, "y": 615, "width": 13, "height": 16, "type": "checkbox"},
    "permits.lifting_plan_permit_no": {"page": 2, "x": 328, "y": 614, "width": 17, "height": 18, "type": "checkbox"},
    "permits.loto_permit_no": {"page": 2, "x": 33, "y": 592, "width": 12, "height": 18, "type": "checkbox"},
    "permits.confined_space_permit_no": {"page": 2, "x": 150, "y": 591, "width": 17, "height": 18, "type": "checkbox"},
    "permits.if_others_text": {"page": 2, "x": 327, "y": 589, "width": 19, "height": 22, "type": "text"},
    # Page 2 — Tool / Equipment (multiline)
    "tools_equipment.list_text": {"page": 2, "x": 32, "y": 492, "width": 549, "height": 62, "type": "multiline"},
    # Page 2 — Permit issuer checklist (tri-state: Y/N/NA)
    "issuer_checks.jsa_carried_out": {"page": 2, "x": 290, "y": 385, "width": 30, "height": 26, "type": "tri_state"},
    "issuer_checks.equipment_deenergized_loto": {"page": 2, "x": 292, "y": 360, "width": 29, "height": 19, "type": "tri_state"},
    "issuer_checks.energized_work_ppe_procedure": {"page": 2, "x": 285, "y": 333, "width": 38, "height": 20, "type": "tri_state"},
    "issuer_checks.workers_fit_trained_qualified": {"page": 2, "x": 292, "y": 307, "width": 33, "height": 22, "type": "tri_state"},
    "issuer_checks.tools_fit_for_voltage": {"page": 2, "x": 292, "y": 307, "width": 33, "height": 22, "type": "tri_state"},
    "issuer_checks.rescue_plan_available": {"page": 2, "x": 287, "y": 252, "width": 38, "height": 24, "type": "tri_state"},
    "issuer_checks.testing_equipment_compatible": {"page": 2, "x": 286, "y": 224, "width": 40, "height": 25, "type": "tri_state"},
    "issuer_checks.line_clearance_taken": {"page": 2, "x": 287, "y": 189, "width": 36, "height": 29, "type": "tri_state"},
    "issuer_checks.environment_condition_suitable": {"page": 2, "x": 546, "y": 386, "width": 36, "height": 21, "type": "tri_state"},
    "issuer_checks.firefighting_equipment_available": {"page": 2, "x": 544, "y": 359, "width": 38, "height": 21, "type": "tri_state"},
    "issuer_checks.rescue_equipment_available": {"page": 2, "x": 547, "y": 333, "width": 33, "height": 20, "type": "tri_state"},
    "issuer_checks.equipment_grounded": {"page": 2, "x": 545, "y": 306, "width": 36, "height": 23, "type": "tri_state"},
    "issuer_checks.adequate_lighting": {"page": 2, "x": 542, "y": 278, "width": 40, "height": 22, "type": "tri_state"},
    "issuer_checks.warning_signage_available": {"page": 2, "x": 542, "y": 253, "width": 40, "height": 20, "type": "tri_state"},
    "issuer_checks.conductive_items_removed": {"page": 2, "x": 544, "y": 224, "width": 37, "height": 24, "type": "tri_state"},
    "issuer_checks.safety_requirements_explained": {"page": 2, "x": 543, "y": 188, "width": 38, "height": 31, "type": "tri_state"},
    # ---------------------------
    # Page 3 — People & signatures (text)
    # ---------------------------
    "people.permit_receiver.name": {"page": 3, "x": 194, "y": 626, "width": 147, "height": 20, "type": "text"},
    "people.permit_receiver.signature": {"page": 3, "x": 350, "y": 626, "width": 110, "height": 18, "type": "text"},
    "people.permit_receiver.datetime": {"page": 3, "x": 468, "y": 626, "width": 109, "height": 17, "type": "text"},
    "people.permit_holder.name": {"page": 3, "x": 191, "y": 602, "width": 150, "height": 20, "type": "text"},
    "people.permit_holder.signature": {"page": 3, "x": 347, "y": 601, "width": 117, "height": 21, "type": "text"},
    "people.permit_holder.datetime": {"page": 3, "x": 468, "y": 600, "width": 110, "height": 23, "type": "text"},
    "people.co_workers.1": {"page": 3, "x": 205, "y": 578, "width": 137, "height": 17, "type": "text"},
    "people.co_workers.2": {"page": 3, "x": 359, "y": 577, "width": 103, "height": 17, "type": "text"},
    "people.co_workers.3": {"page": 3, "x": 479, "y": 578, "width": 101, "height": 16, "type": "text"},
    "people.co_workers.4": {"page": 3, "x": 202, "y": 554, "width": 142, "height": 20, "type": "text"},
    "people.co_workers.5": {"page": 3, "x": 358, "y": 554, "width": 103, "height": 19, "type": "text"},
    "people.co_workers.6": {"page": 3, "x": 475, "y": 554, "width": 105, "height": 19, "type": "text"},
    "people.permit_issuer.name": {"page": 3, "x": 193, "y": 509, "width": 148, "height": 17, "type": "text"},
    "people.permit_issuer.signature": {"page": 3, "x": 349, "y": 508, "width": 112, "height": 18, "type": "text"},
    "people.permit_issuer.datetime": {"page": 3, "x": 466, "y": 508, "width": 113, "height": 17, "type": "text"},
    # Page 3 — Extension (text + multiline)
    "extension.date": {"page": 3, "x": 28, "y": 415, "width": 76, "height": 19, "type": "text"},
    "extension.time_from": {"page": 3, "x": 107, "y": 416, "width": 31, "height": 15, "type": "text"},
    "extension.time_to": {"page": 3, "x": 143, "y": 416, "width": 38, "height": 15, "type": "text"},
    "extension.permit_holder.name": {"page": 3, "x": 187, "y": 414, "width": 77, "height": 17, "type": "text"},
    "extension.permit_holder.signature": {"page": 3, "x": 271, "y": 415, "width": 45, "height": 16, "type": "text"},
    "extension.permit_receiver.name": {"page": 3, "x": 320, "y": 414, "width": 75, "height": 15, "type": "text"},
    "extension.permit_receiver.signature": {"page": 3, "x": 396, "y": 416, "width": 48, "height": 14, "type": "text"},
    "extension.permit_issuer.name": {"page": 3, "x": 447, "y": 413, "width": 84, "height": 18, "type": "text"},
    "extension.permit_issuer.signature": {"page": 3, "x": 533, "y": 415, "width": 45, "height": 15, "type": "text"},
    "extension.remarks": {"page": 3, "x": 96, "y": 372, "width": 479, "height": 36, "type": "multiline"},
    # Page 3 — Closure (text)
    "closure.permit_receiver.name": {"page": 3, "x": 191, "y": 267, "width": 149, "height": 20, "type": "text"},
    "closure.permit_receiver.signature": {"page": 3, "x": 348, "y": 266, "width": 116, "height": 21, "type": "text"},
    "closure.permit_receiver.datetime": {"page": 3, "x": 468, "y": 266, "width": 111, "height": 20, "type": "text"},
    "closure.permit_holder.name": {"page": 3, "x": 194, "y": 244, "width": 148, "height": 17, "type": "text"},
    "closure.permit_holder.signature": {"page": 3, "x": 349, "y": 244, "width": 113, "height": 15, "type": "text"},
    "closure.permit_holder.datetime": {"page": 3, "x": 466, "y": 244, "width": 113, "height": 17, "type": "text"},
    "closure.permit_issuer.name": {"page": 3, "x": 191, "y": 220, "width": 153, "height": 18, "type": "text"},
    "closure.permit_issuer.signature": {"page": 3, "x": 348, "y": 221, "width": 112, "height": 18, "type": "text"},
    "closure.permit_issuer.datetime": {"page": 3, "x": 466, "y": 220, "width": 114, "height": 19, "type": "text"},
}

# Read-only mapping from overlay tags to existing form_data keys (keeps DB/UI schema unchanged).
_PTW_TAG_SOURCES: dict[str, list[str]] = {
    # Meta
    "meta.permit_no": ["permit_no", "work_order_id"],
    "meta.permit_validity_date": ["permit_validity_date"],
    "meta.start_time": ["start_time"],
    "meta.end_time": ["end_time"],
    "meta.project_site_name": ["project_site_name", "site_name"],
    "meta.work_location": ["work_location"],
    "meta.description_of_work": ["description_of_work", "work_description", "work_description_line1"],
    "meta.contractor_name": ["contractor_name"],
    # Hazards
    "hazards.live_dc_cables": ["hz_live_dc_cables"],
    "hazards.loose_or_damaged_connectors": ["hz_loose_connectors"],
    "hazards.moving_solar_tracker_parts": ["hz_tracker_parts"],
    "hazards.dust_accumulation_on_equipment": ["hz_dust"],
    "hazards.high_dc_voltage_from_pv_panels": ["hz_high_dc"],
    "hazards.poor_grounding_earthing": ["hz_poor_grounding"],
    "hazards.heavy_solar_panels": ["hz_heavy_panels"],
    "hazards.wildlife_snakes_insects_birds": ["hz_wildlife"],
    "hazards.arc_flash_short_circuit": ["hz_arc_flash"],
    "hazards.working_at_height_rooftop_structure": ["hz_working_height"],
    "hazards.sharp_panel_edges_metal_structures": ["hz_sharp_edges"],
    "hazards.lightning_or_stormy_weather": ["hz_lightning"],
    "hazards.improper_grounding": ["hz_improper_grounding"],
    "hazards.wet_surfaces": ["hz_wet_surfaces"],
    "hazards.high_ambient_temperature_sun_exposure": ["hz_heat"],
    "hazards.overloaded_circuits": ["hz_overload"],
    "hazards.manual_handling": ["hz_manual_handling"],
    "hazards.overhead_line": ["hz_overhead_line"],
    "hazards.if_others_text": ["hz_others_text"],
    # Risks
    "risks.electrocution": ["rk_electrocution"],
    "risks.burns_equipment_damage": ["rk_burns"],
    "risks.unexpected_energization": ["rk_unexpected_energization"],
    "risks.heat_stress_dehydration": ["rk_heat_stress"],
    "risks.electric_shock": ["rk_electric_shock"],
    "risks.fire_or_overheating": ["rk_fire"],
    "risks.crushing_or_pinching_injuries": ["rk_crushing"],
    "risks.electric_burn": ["rk_electric_burn"],
    "risks.fall_causing_serious_injury": ["rk_fall"],
    "risks.back_or_muscle_injury": ["rk_back_injury"],
    "risks.bites_stings": ["rk_bites"],
    "risks.falling_particles": ["rk_falling_particles"],
    "risks.tripping_slipping": ["rk_tripping"],
    "risks.if_others_text": ["rk_others_text"],
    # PPE
    "ppe.safety_helmet": ["ppe_helmet"],
    "ppe.hrc_suit": ["ppe_hrc_suit"],
    "ppe.respirators": ["ppe_respirators", "ppe_respirator"],
    "ppe.full_body_safety_harness": ["ppe_harness"],
    "ppe.safety_shoes": ["ppe_shoes"],
    "ppe.electrical_mat": ["ppe_electrical_mat"],
    "ppe.dust_masks": ["ppe_dust_mask"],
    "ppe.lifeline_anchor_point": ["ppe_lifeline"],
    "ppe.reflective_vest": ["ppe_reflective_vest"],
    "ppe.face_shield_with_arc_protection": ["ppe_face_shield"],
    "ppe.ear_plugs_muffs": ["ppe_ear_plugs"],
    "ppe.cut_resistant_gloves": ["ppe_cut_gloves"],
    "ppe.safety_goggles": ["ppe_goggles"],
    "ppe.insulated_hand_tools": ["ppe_insulated_tools"],
    "ppe.electrical_hand_gloves": ["ppe_electrical_gloves"],
    "ppe.if_others_text": ["ppe_others_text"],
    # Precautions
    "precautions.electrical_isolation": ["sp_electrical_isolation"],
    "precautions.fire_extinguisher_first_aid": ["sp_fire_extinguisher"],
    "precautions.proper_isolation": ["sp_proper_isolation"],
    "precautions.authorized_competent_personnel": ["sp_authorized_personnel"],
    "precautions.lockout_tagout": ["sp_loto"],
    "precautions.warning_signage_barricading": ["sp_signage"],
    "precautions.rescue_equipment_at_site": ["sp_rescue_equipment"],
    "precautions.zero_voltage_ensure": ["sp_zero_voltage"],
    "precautions.proper_earthing_grounding": ["sp_earthing"],
    "precautions.pre_job_meeting": ["sp_pre_job_meeting"],
    "precautions.insulated_tools": ["sp_insulated_tools"],
    "precautions.proper_illumination": ["sp_illumination"],
    "precautions.escape_route_clear": ["sp_escape_route"],
    "precautions.if_others_text": ["sp_others_text"],
    # Associated permits / documents (S1 stores booleans only)
    "permits.hot_work_permit_no": ["ap_hot_work"],
    "permits.work_at_night_permit_no": ["ap_night_work"],
    "permits.work_at_height_permit_no": ["ap_height_work"],
    "permits.general_work_permit_no": ["ap_general_work"],
    "permits.excavation_work_permit_no": ["ap_excavation"],
    "permits.lifting_plan_permit_no": ["ap_lifting"],
    "permits.loto_permit_no": ["ap_loto"],
    "permits.confined_space_permit_no": ["ap_confined_space"],
    "permits.if_others_text": ["ap_others_text"],
    # Tools / equipment
    "tools_equipment.list_text": ["tools_equipment"],
    # Issuer checks (tri-state values: "Y"/"N"/"NA")
    "issuer_checks.jsa_carried_out": ["chk_jsa"],
    "issuer_checks.environment_condition_suitable": ["chk_environment"],
    "issuer_checks.equipment_deenergized_loto": ["chk_loto"],
    "issuer_checks.firefighting_equipment_available": ["chk_fire_fighting"],
    "issuer_checks.energized_work_ppe_procedure": ["chk_energized_ppe"],
    "issuer_checks.rescue_equipment_available": ["chk_rescue"],
    "issuer_checks.workers_fit_trained_qualified": ["chk_workers_fit"],
    "issuer_checks.equipment_grounded": ["chk_grounded"],
    "issuer_checks.tools_fit_for_voltage": ["chk_tools"],
    "issuer_checks.adequate_lighting": ["chk_lighting"],
    "issuer_checks.rescue_plan_available": ["chk_rescue_plan"],
    "issuer_checks.warning_signage_available": ["chk_signage"],
    "issuer_checks.testing_equipment_compatible": ["chk_testing_equipment"],
    "issuer_checks.conductive_items_removed": ["chk_conductive_removed"],
    "issuer_checks.line_clearance_taken": ["chk_line_clearance"],
    "issuer_checks.safety_requirements_explained": ["chk_briefing"],
    # People (names/dates are stored across S1/S2/S3)
    "people.permit_receiver.name": ["receiver_name"],
    "people.permit_receiver.signature": ["receiver_signature"],
    "people.permit_receiver.datetime": ["receiver_datetime"],
    "people.permit_holder.name": ["permit_holder_name", "holder_name"],
    "people.permit_holder.signature": ["holder_signature"],
    "people.permit_holder.datetime": ["permit_holder_datetime", "holder_datetime"],
    "people.permit_issuer.name": ["permit_issuer_name", "issuer_name"],
    "people.permit_issuer.signature": ["issuer_signature"],
    "people.permit_issuer.datetime": ["permit_issuer_datetime", "issuer_datetime"],
    "people.co_workers.1": ["coworker_1"],
    "people.co_workers.2": ["coworker_2"],
    "people.co_workers.3": ["coworker_3"],
    "people.co_workers.4": ["coworker_4"],
    "people.co_workers.5": ["coworker_5"],
    "people.co_workers.6": ["coworker_6"],
    # Extension / Closure (legacy flat keys from earlier overlay)
    "extension.date": ["ext_date"],
    "extension.time_from": ["ext_from_time"],
    "extension.time_to": ["ext_to_time"],
    "extension.permit_holder.name": ["ext_holder_name"],
    "extension.permit_holder.signature": ["ext_holder_signature"],
    "extension.permit_receiver.name": ["ext_receiver_name"],
    "extension.permit_receiver.signature": ["ext_receiver_signature"],
    "extension.permit_issuer.name": ["ext_issuer_name"],
    "extension.permit_issuer.signature": ["ext_issuer_signature"],
    "extension.remarks": ["ext_remarks"],
    "closure.permit_receiver.name": ["closure_receiver"],
    "closure.permit_receiver.signature": ["closure_receiver_signature"],
    "closure.permit_receiver.datetime": ["closure_receiver_datetime"],
    "closure.permit_holder.name": ["closure_holder"],
    "closure.permit_holder.signature": ["closure_holder_signature"],
    "closure.permit_holder.datetime": ["closure_holder_datetime"],
    "closure.permit_issuer.name": ["closure_issuer"],
    "closure.permit_issuer.signature": ["closure_issuer_signature"],
    "closure.permit_issuer.datetime": ["closure_issuer_datetime"],
}

# Approved stamp position (kept identical to previous behavior; offsets were baked in)
_APPROVED_STAMP_X = 374
_APPROVED_STAMP_Y = 790
_APPROVED_STAMP_DATE_Y = 772


def _draw_approved_stamp(canvas_obj: Any, approved_datetime: Any) -> None:
    """
    Draw APPROVED stamp (page 1 only).

    Condition is handled by caller.
    Data source: form_data["date_s3_approved"]
    """
    stamp_x = _APPROVED_STAMP_X
    stamp_y = _APPROVED_STAMP_Y

    # "APPROVED" in green
    canvas_obj.setFont("Helvetica-Bold", 16)
    canvas_obj.setFillColorRGB(0, 0.6, 0)
    canvas_obj.drawString(stamp_x, stamp_y, "APPROVED")

    # Date below, in black
    canvas_obj.setFont("Helvetica", 9)
    canvas_obj.setFillColorRGB(0, 0, 0)
    canvas_obj.drawString(
        stamp_x,
        _APPROVED_STAMP_DATE_Y,
        f"Date: {'' if approved_datetime is None else str(approved_datetime)}",
    )


@st.cache_data(show_spinner=False, ttl=3600)
def _download_pdf_template_from_supabase() -> bytes:
    """
    Download the PDF template specifically for ReportLab overlay approach.
    Returns the PDF template bytes; raises RuntimeError if not found.
    """
    sb = get_supabase_client(prefer_service_role=True)
    try:
        response = sb.storage.from_(SUPABASE_TEMPLATE_BUCKET).download(TEMPLATE_PDF_FILE_NAME)
        if isinstance(response, (bytes, bytearray)) and len(response) > 0:
            if bytes(response)[:4] == b"%PDF":
                return bytes(response)
            raise RuntimeError(f"Downloaded file is not a valid PDF: {TEMPLATE_PDF_FILE_NAME}")
        raise RuntimeError(f"Empty PDF template received: {TEMPLATE_PDF_FILE_NAME}")
    except Exception as e:
        raise RuntimeError(
            f"Failed to download PDF template from Supabase.\n"
            f"Bucket: {SUPABASE_TEMPLATE_BUCKET}\n"
            f"File: {TEMPLATE_PDF_FILE_NAME}\n"
            f"Error: {e}\n\n"
            f"Please upload '{TEMPLATE_PDF_FILE_NAME}' to the bucket."
        ) from e


def _is_truthy(val: Any) -> bool:
    """Check if a value is truthy for checkbox rendering."""
    if val is None:
        return False
    if isinstance(val, bool):
        return val
    if isinstance(val, str):
        return val.strip().upper() in {"Y", "YES", "TRUE", "1", "✓", "✔"}
    return bool(val)


def _prepare_overlay_data(form_data: dict) -> dict:
    """
    Prepare a read-only overlay data dict keyed by PTW_PDF_COORDINATES dotted tags.

    IMPORTANT:
    - Does NOT mutate or change the DB/UI schema.
    - Uses _PTW_TAG_SOURCES to map existing flat form_data into dotted overlay tags.
    """
    fd = form_data if isinstance(form_data, dict) else {}
    out: dict[str, Any] = {}

    for tag in PTW_PDF_COORDINATES.keys():
        sources = _PTW_TAG_SOURCES.get(tag, [])
        val = None
        for k in sources:
            if k in fd:
                val = fd.get(k)
                # Keep explicit False/0 for checkbox-ish fields; empties get filtered at render-time
                break
        out[tag] = val

    return out


def _ptw_wrap_text_to_width(text: str, *, font_name: str, font_size: int, max_width: float) -> list[str]:
    """Wrap text into lines that fit max_width using ReportLab string metrics."""
    try:
        from reportlab.pdfbase import pdfmetrics  # type: ignore
    except Exception:
        # Fallback: crude wrap by characters
        import textwrap
        return textwrap.wrap(text, width=max(1, int(max_width // 5))) or [text]

    words = (text or "").replace("\r", " ").split()
    if not words:
        return []

    lines: list[str] = []
    cur = words[0]
    for w in words[1:]:
        candidate = f"{cur} {w}"
        if pdfmetrics.stringWidth(candidate, font_name, font_size) <= max_width:
            cur = candidate
        else:
            lines.append(cur)
            cur = w
    lines.append(cur)
    return lines


def _ptw_draw_centered_text(
    canvas_obj: Any,
    *,
    x: float,
    y: float,
    w: float,
    h: float,
    text: str,
    font_name: str = "Helvetica-Bold",
    font_size: int = 10,
) -> None:
    """Draw a short text centered in the given rectangle."""
    try:
        from reportlab.pdfbase import pdfmetrics  # type: ignore
        tw = float(pdfmetrics.stringWidth(text, font_name, font_size))
    except Exception:
        tw = float(len(text) * (font_size * 0.6))
    tx = float(x) + max(0.0, (float(w) - tw) / 2.0)
    ty = float(y) + max(0.0, (float(h) - float(font_size)) / 2.0)
    canvas_obj.setFont(font_name, font_size)
    canvas_obj.drawString(tx, ty, text)


def _ptw_draw_text_in_box(
    canvas_obj: Any,
    *,
    x: float,
    y: float,
    w: float,
    h: float,
    text: str,
    multiline: bool,
    font_name: str = "Helvetica",
    font_size: int = 9,
) -> None:
    """Draw text inside a rectangular box; clips/wraps to fit."""
    s = ("" if text is None else str(text)).replace("\r\n", "\n").replace("\r", "\n").strip()
    if not s:
        return

    canvas_obj.setFont(font_name, font_size)

    if multiline:
        lines = []
        for para in s.split("\n"):
            para = para.strip()
            if not para:
                lines.append("")
                continue
            lines.extend(_ptw_wrap_text_to_width(para, font_name=font_name, font_size=font_size, max_width=float(w)))

        line_h = float(font_size) + 1.0
        max_lines = max(1, int(float(h) // line_h))
        lines = lines[:max_lines]

        # Draw from top down within the box
        y_top = float(y) + float(h) - float(font_size)
        for i, line in enumerate(lines):
            yy = y_top - (i * line_h)
            if yy < float(y):
                break
            canvas_obj.drawString(float(x), yy, line)
        return

    # Single-line: clip with ellipsis if needed
    try:
        from reportlab.pdfbase import pdfmetrics  # type: ignore

        if pdfmetrics.stringWidth(s, font_name, font_size) > float(w):
            ell = "…"
            # Ensure ellipsis fits
            while s and pdfmetrics.stringWidth(s + ell, font_name, font_size) > float(w):
                s = s[:-1]
            s = (s + ell) if s else ""
    except Exception:
        # Crude truncation fallback
        max_chars = max(1, int(float(w) // (font_size * 0.6)))
        if len(s) > max_chars:
            s = s[: max(1, max_chars - 1)] + "…"

    yy = float(y) + max(1.0, (float(h) - float(font_size)) / 2.0)
    canvas_obj.drawString(float(x), yy, s)


def _render_ptw_overlay_page(canvas_obj: Any, *, page_num: int, data: dict) -> None:
    """
    Render overlay content for a single page using PTW_PDF_COORDINATES.
    """
    for tag, cfg in PTW_PDF_COORDINATES.items():
        if int(cfg.get("page", 0)) != int(page_num):
            continue

        v = (data or {}).get(tag)
        if v is None:
            continue

        x = float(cfg.get("x", 0))
        y = float(cfg.get("y", 0))
        w = float(cfg.get("width", 0))
        h = float(cfg.get("height", 0))
        typ = str(cfg.get("type") or "text").strip().lower()

        if typ == "checkbox":
            if _is_truthy(v):
                # Center a checkmark in the provided rectangle
                fs = max(8, min(14, int(h) + 2))
                _ptw_draw_centered_text(canvas_obj, x=x, y=y, w=w, h=h, text="✓", font_name="Helvetica-Bold", font_size=fs)
            continue

        if typ == "tri_state":
            s = str(v).strip().upper()
            if s in {"Y", "N", "NA"}:
                _ptw_draw_centered_text(canvas_obj, x=x, y=y, w=w, h=h, text=s, font_name="Helvetica-Bold", font_size=9)
            continue

        if typ == "multiline":
            _ptw_draw_text_in_box(canvas_obj, x=x, y=y, w=w, h=h, text=str(v), multiline=True, font_name="Helvetica", font_size=8)
            continue

        # default: text
        _ptw_draw_text_in_box(canvas_obj, x=x, y=y, w=w, h=h, text=str(v), multiline=False, font_name="Helvetica", font_size=9)


    return


def generate_ptw_pdf_from_template(form_data: dict, *, progress_callback=None) -> bytes:
    """
    Generate PTW PDF using PDF template + ReportLab overlay.
    
    This approach:
    - Downloads the PDF template from Supabase
    - Creates an overlay with text and checkmarks using ReportLab
    - Merges the overlay with the template using PyPDF2
    
    NO DOCX. NO LibreOffice. NO paid services.
    Works on Streamlit Cloud (Linux, no system binaries).
    
    Args:
        form_data: Dictionary of PTW form field values
        progress_callback: Optional callback(percent, message)
    
    Returns:
        Final PDF bytes
    
    Raises:
        RuntimeError if generation fails
    """
    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.pdfgen import canvas as rl_canvas  # type: ignore
        from PyPDF2 import PdfReader, PdfWriter  # type: ignore
    except ImportError as e:
        raise RuntimeError(f"Required PDF libraries not installed: {e}") from e
    
    if progress_callback:
        try:
            progress_callback(5, "Downloading PDF template...")
        except Exception:
            pass
    
    # Download PDF template
    template_bytes = _download_pdf_template_from_supabase()
    
    if progress_callback:
        try:
            progress_callback(20, "Preparing form data...")
        except Exception:
            pass
    
    # Prepare data for overlay
    data = _prepare_overlay_data(form_data)
    
    # Read template to get number of pages
    template_reader = PdfReader(BytesIO(template_bytes))
    num_pages = len(template_reader.pages)
    
    if progress_callback:
        try:
            progress_callback(30, "Creating overlay...")
        except Exception:
            pass
    
    # Create overlay PDF with ReportLab (one page per template page)
    overlay_buffer = BytesIO()
    c = rl_canvas.Canvas(overlay_buffer, pagesize=A4)

    for page_idx in range(num_pages):
        _render_ptw_overlay_page(c, page_num=page_idx + 1, data=data)

        # APPROVED stamp (page 1 only), in the PDF overlay pipeline
        if (
            page_idx == 0
            and str(form_data.get("status", "")).strip().upper() == "APPROVED"
        ):
            _draw_approved_stamp(c, form_data.get("date_s3_approved"))

        c.showPage()
    
    c.save()
    overlay_buffer.seek(0)
    
    if progress_callback:
        try:
            progress_callback(60, "Merging PDF layers...")
        except Exception:
            pass
    
    # Read overlay PDF
    overlay_reader = PdfReader(overlay_buffer)
    
    # Merge: for each page, merge template page with overlay page
    writer = PdfWriter()
    for page_idx in range(num_pages):
        template_page = template_reader.pages[page_idx]
        
        if page_idx < len(overlay_reader.pages):
            overlay_page = overlay_reader.pages[page_idx]
            # Merge overlay on top of template
            template_page.merge_page(overlay_page)
        
        writer.add_page(template_page)
    
    if progress_callback:
        try:
            progress_callback(90, "Finalizing PDF...")
        except Exception:
            pass
    
    # Write final PDF
    output_buffer = BytesIO()
    writer.write(output_buffer)
    output_buffer.seek(0)
    pdf_bytes = output_buffer.read()
    
    # Validate output
    if not pdf_bytes or pdf_bytes[:4] != b"%PDF":
        raise RuntimeError("PDF generation failed: output is not a valid PDF")
    
    if progress_callback:
        try:
            progress_callback(100, "PDF ready!")
        except Exception:
            pass
    
    return pdf_bytes


def _add_overlay_elements(pdf_bytes: bytes, form_data: dict) -> bytes:
    """
    Post-process overlay for specific elements that require box-aware wrapping / verification.
    - Undertaking checkbox (page 2)
    - Tools/Equipment Required block (page 2)
    """
    try:
        from reportlab.pdfgen import canvas  # type: ignore
        from reportlab.lib.pagesizes import A4  # type: ignore
        from PyPDF2 import PdfReader, PdfWriter  # type: ignore
    except Exception:
        return pdf_bytes

    if not isinstance(pdf_bytes, (bytes, bytearray)) or bytes(pdf_bytes)[:4] != b"%PDF":
        return pdf_bytes

    try:
        reader = PdfReader(BytesIO(bytes(pdf_bytes)))
        writer = PdfWriter()

        packet = BytesIO()
        c = canvas.Canvas(packet, pagesize=A4)

        # ---- Undertaking Checkbox ----
        if bool((form_data or {}).get("undertaking_checkbox")):
            c.setFont("Helvetica-Bold", 14)
            c.drawString(209, 91, "✓")

        # ---- Tools/Equipment Required ----
        tools_text = str((form_data or {}).get("tools_equipment_required") or "").strip()
        if tools_text:
            x = 32
            y = 491
            max_width = 550
            line_height = 12

            lines: list[str] = []
            words = tools_text.split()
            current = ""

            for word in words:
                test_line = (current + " " + word).strip() if current else word
                if c.stringWidth(test_line, "Helvetica", 10) <= max_width:
                    current = test_line
                else:
                    if current:
                        lines.append(current)
                    current = word
            if current:
                lines.append(current)

            c.setFont("Helvetica", 10)
            for line in lines[:5]:  # limit height
                c.drawString(x, y, line)
                y -= line_height

        c.save()
        packet.seek(0)
        overlay_pdf = PdfReader(packet)

        for i in range(len(reader.pages)):
            page = reader.pages[i]
            if i == 1 and overlay_pdf.pages:  # page 2 (0-based index)
                page.merge_page(overlay_pdf.pages[0])
            writer.add_page(page)

        output = BytesIO()
        writer.write(output)
        output.seek(0)
        out = output.read()
        return out if out[:4] == b"%PDF" else bytes(pdf_bytes)
    except Exception:
        return bytes(pdf_bytes)


def _fill_pdf_template_acroform(pdf_template_bytes: bytes, data: dict) -> bytes | None:
    """
    Fill a PDF template that contains AcroForm fields.

    - If the template has no AcroForm fields, returns None.
    - Uses PyPDF2 to set field values and enables NeedAppearances.
    """
    try:
        from PyPDF2 import PdfReader, PdfWriter  # type: ignore
    except Exception:
        return None

    try:
        reader = PdfReader(BytesIO(pdf_template_bytes))
        fields = reader.get_fields() or {}
        if not fields:
            return None

        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)

        # Ensure appearances are generated in most PDF viewers
        try:
            if "/AcroForm" in reader.trailer["/Root"]:
                acro = dict(reader.trailer["/Root"]["/AcroForm"])
            else:
                acro = {}
            acro.update({"/NeedAppearances": True})
            writer._root_object.update({"/AcroForm": acro})  # type: ignore[attr-defined]
        except Exception:
            pass

        # Only set values for fields that exist in the template
        values: dict[str, Any] = {}
        for name, field in fields.items():
            if name not in data:
                continue
            v = data.get(name)
            ft = None
            try:
                ft = field.get("/FT")
            except Exception:
                ft = None
            if ft == "/Btn":
                # Checkbox/radio buttons: map common tick marks / truthy strings to "Yes"
                s = "" if v is None else str(v).strip()
                values[name] = "Yes" if s in {"✓", "✔", "Y", "YES", "TRUE", "1"} else "Off"
            else:
                values[name] = "" if v is None else str(v)

        for page in writer.pages:
            try:
                writer.update_page_form_field_values(page, values)  # type: ignore[arg-type]
            except Exception:
                pass

        out = BytesIO()
        writer.write(out)
        out.seek(0)
        pdf_bytes = out.read()
        return pdf_bytes if pdf_bytes[:4] == b"%PDF" else None
    except Exception:
        return None


def convert_docx_to_pdf(docx_bytes: bytes, *, progress_callback=None) -> bytes:
    """
    Convert DOCX bytes to PDF bytes using MS Word (pywin32).
    
    For legal documents like PTW, PDF is mandatory.
    
    Returns:
        PDF bytes if conversion succeeds.
        Original DOCX bytes if conversion fails (caller should check).
    """
    import logging
    import shutil
    logger = logging.getLogger(__name__)
    
    # Create temp files in a simple path (avoid special characters)
    temp_dir = Path(os.environ.get("TEMP", os.environ.get("TMP", ".")))
    docx_path = temp_dir / f"ptw_temp_{os.getpid()}.docx"
    pdf_path = temp_dir / f"ptw_temp_{os.getpid()}.pdf"
    
    try:
        # Write DOCX to temp file
        docx_path.write_bytes(docx_bytes)
        
        # Mode: cloud-first / local-first (default) / cloud-only
        mode = (_get_setting("PTW_PDF_CONVERTER_MODE") or "auto").strip().lower()
        cloud_first = mode in {"cloud", "cloud_first", "cloud-first"}
        cloud_only = mode in {"cloud_only", "cloud-only"}

        # Cloud-first (robust for hosted deployments)
        if cloud_first or cloud_only:
            cloud_pdf = _convert_docx_to_pdf_cloudconvert(docx_bytes, progress_callback=progress_callback)
            if cloud_pdf is not None:
                return cloud_pdf
            if cloud_only:
                return docx_bytes  # caller will raise with compliance message

        # Method 1: Try pywin32 (win32com.client) - most reliable on Windows
        if os.name == 'nt':
            try:
                import win32com.client  # pyright: ignore[reportMissingImports]
                import pythoncom  # pyright: ignore[reportMissingImports]
                
                # Initialize COM for this thread
                pythoncom.CoInitialize()
                
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    word.DisplayAlerts = False
                    
                    # Use absolute paths with forward slashes for COM
                    docx_abs = str(docx_path.resolve()).replace("/", "\\")
                    pdf_abs = str(pdf_path.resolve()).replace("/", "\\")
                    
                    doc = word.Documents.Open(docx_abs, ReadOnly=True)
                    # wdFormatPDF = 17
                    doc.SaveAs2(pdf_abs, FileFormat=17)
                    doc.Close(SaveChanges=False)
                    word.Quit()
                    
                    if pdf_path.exists() and pdf_path.stat().st_size > 0:
                        pdf_bytes = pdf_path.read_bytes()
                        return pdf_bytes
                        
                except Exception as e:
                    logger.warning(f"pywin32 Word conversion failed: {e}")
                finally:
                    pythoncom.CoUninitialize()
                    
            except ImportError:
                logger.warning("pywin32 not available")
        
        # Method 2: Try docx2pdf library
        try:
            from docx2pdf import convert as docx2pdf_convert  # pyright: ignore[reportMissingImports]
            docx2pdf_convert(str(docx_path), str(pdf_path))
            if pdf_path.exists() and pdf_path.stat().st_size > 0:
                return pdf_path.read_bytes()
        except ImportError:
            logger.warning("docx2pdf not available")
        except Exception as e:
            logger.warning(f"docx2pdf conversion failed: {e}")
        
        # Method 3: Try LibreOffice
        soffice_paths: list[str] = []
        # Prefer PATH-based discovery (works on Linux/Cloud/Mac when LibreOffice is installed)
        for cmd in ("soffice", "libreoffice"):
            p = shutil.which(cmd)
            if p:
                soffice_paths.append(p)
        # Common install locations
        soffice_paths.extend(
            [
                # Windows
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                # Linux
                "/usr/bin/soffice",
                "/usr/bin/libreoffice",
                "/snap/bin/libreoffice",
                "/usr/lib/libreoffice/program/soffice",
                "/usr/lib64/libreoffice/program/soffice",
                # macOS
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            ]
        )
        # De-dup while preserving order
        seen = set()
        soffice_paths = [p for p in soffice_paths if p and (p not in seen and not seen.add(p))]

        for sp in soffice_paths:
            if not Path(sp).exists():
                continue
            try:
                # LibreOffice writes <basename>.pdf into outdir
                subprocess.run(
                    [
                        sp,
                        "--headless",
                        "--nologo",
                        "--nolockcheck",
                        "--norestore",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(temp_dir),
                        str(docx_path),
                    ],
                    capture_output=True,
                    timeout=120,
                )
                if pdf_path.exists() and pdf_path.stat().st_size > 0:
                    return pdf_path.read_bytes()
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed ({sp}): {e}")

        # Local-first fallback to CloudConvert (works on hosted deployments)
        if not cloud_only:
            cloud_pdf = _convert_docx_to_pdf_cloudconvert(docx_bytes, progress_callback=progress_callback)
            if cloud_pdf is not None:
                return cloud_pdf
                
    finally:
        # Clean up temp files
        try:
            if docx_path.exists():
                docx_path.unlink()
            if pdf_path.exists():
                pdf_path.unlink()
        except Exception:
            pass
    
    # Return original DOCX if all methods fail
    return docx_bytes


# =============================================================================
# S3 Approval Timestamp + APPROVED Stamp (for S1 downloads)
# =============================================================================


def _s1_get_ptw_approval_times(work_order_id: str) -> dict:
    """
    Fetch S2 and S3 approval timestamps from work_orders (single source of truth).

    Returns:
        {
          "holder_datetime": "DD-MM-YYYY HH:MM" (from date_s2_forwarded),
          "issuer_datetime": "DD-MM-YYYY HH:MM" (from date_s3_approved),
          "date_s2_forwarded_raw": original timestamp or None,
          "date_s3_approved_raw": original timestamp or None,
        }
    """
    sb = get_supabase_client(prefer_service_role=True)
    resp = (
        sb.table(TABLE_WORK_ORDERS)
        .select("date_s2_forwarded,date_s3_approved")
        .eq("work_order_id", work_order_id)
        .limit(1)
        .execute()
    )
    err = getattr(resp, "error", None)
    if err:
        raise RuntimeError(f"Failed to fetch approval times: {err}")
    rows = getattr(resp, "data", None) or []
    if not rows:
        return {
            "holder_datetime": "",
            "issuer_datetime": "",
            "date_s2_forwarded_raw": None,
            "date_s3_approved_raw": None,
        }

    row = rows[0]
    s2_raw = row.get("date_s2_forwarded")
    s3_raw = row.get("date_s3_approved")

    def _fmt(val) -> str:
        if val is None or str(val).strip() == "":
            return ""
        try:
            return pd.to_datetime(val).strftime("%d-%m-%Y %H:%M")
        except Exception:
            return str(val)

    return {
        "holder_datetime": _fmt(s2_raw),
        "issuer_datetime": _fmt(s3_raw),
        "date_s2_forwarded_raw": s2_raw,
        "date_s3_approved_raw": s3_raw,
    }


def _s1_add_floating_approval_stamp(pdf_bytes: bytes, *, approved_on: str) -> bytes:
    """
    Add floating APPROVED stamp overlay to EVERY page (post S3 approval).
    Does not modify DOCX/template; PDF-only overlay.
    """
    if not _HAS_S1_PDF_STAMP_LIBS:
        return pdf_bytes

    try:
        # Build stamp overlay
        stamp_buffer = BytesIO()
        c = canvas.Canvas(stamp_buffer, pagesize=A4)
        page_width, _page_height = A4

        # Position near Permit Issuer section (tunable)
        stamp_x = page_width - 200
        stamp_y = 120
        stamp_width = 160
        stamp_height = 60

        stamp_color = Color(0.8, 0.1, 0.1, alpha=0.85)
        c.setStrokeColor(stamp_color)
        c.setLineWidth(3)
        c.rect(stamp_x, stamp_y, stamp_width, stamp_height, stroke=1, fill=0)
        c.setLineWidth(1.5)
        c.rect(stamp_x + 4, stamp_y + 4, stamp_width - 8, stamp_height - 8, stroke=1, fill=0)

        c.setFillColor(stamp_color)
        c.setFont("Helvetica-Bold", 18)
        center_x = stamp_x + stamp_width / 2
        c.drawCentredString(center_x, stamp_y + 35, "APPROVED")
        c.setFont("Helvetica", 9)
        if approved_on:
            c.drawCentredString(center_x, stamp_y + 18, f"ON: {approved_on}")

        c.save()
        stamp_buffer.seek(0)

        original_pdf = PdfReader(BytesIO(pdf_bytes))
        stamp_pdf = PdfReader(stamp_buffer)
        stamp_page = stamp_pdf.pages[0]

        out = PdfWriter()
        for i in range(len(original_pdf.pages)):
            p = original_pdf.pages[i]
            p.merge_page(stamp_page)
            out.add_page(p)

        out_buf = BytesIO()
        out.write(out_buf)
        out_buf.seek(0)
        return out_buf.read()
    except Exception:
        # Never break download due to stamp issues
        return pdf_bytes


def _s1_inject_approval_times(form_data: dict, work_order_id: str) -> tuple[dict, dict]:
    """Inject holder_datetime / issuer_datetime from work_orders into form_data."""
    approval_times = _s1_get_ptw_approval_times(work_order_id)
    updated = dict(form_data) if isinstance(form_data, dict) else {}
    if approval_times.get("holder_datetime"):
        updated["holder_datetime"] = approval_times["holder_datetime"]
    if approval_times.get("issuer_datetime"):
        updated["issuer_datetime"] = approval_times["issuer_datetime"]
    return updated, approval_times


def generate_ptw_pdf(template_bytes: bytes, form_data: dict, *, progress_callback=None) -> bytes:
    """
    Generate PTW document as PDF (PDF-only).
    
    PRIMARY METHOD: PDF Template + ReportLab overlay (no DOCX/LibreOffice/CloudConvert)
    FALLBACK: DOCX conversion (local MS Word/LibreOffice or CloudConvert)
    
    Args:
        template_bytes: The template file as bytes (ignored for primary method)
        form_data: Dictionary of placeholder keys to values
        progress_callback: Optional callback(percent, message)
    
    Returns:
        PDF bytes.

    Raises:
        RuntimeError if PDF generation failed.
    """
    # -------------------------------------------------------------------------
    # PRIMARY METHOD: PDF Template + ReportLab Overlay
    # Works on Streamlit Cloud (Linux) without any system binaries or paid services
    # -------------------------------------------------------------------------
    try:
        if progress_callback:
            try:
                progress_callback(5, "Generating PDF from template...")
            except Exception:
                pass
        
        pdf_bytes = generate_ptw_pdf_from_template(form_data, progress_callback=progress_callback)
        
        if isinstance(pdf_bytes, (bytes, bytearray)) and pdf_bytes[:4] == b"%PDF":
            return bytes(pdf_bytes)
    except Exception as template_err:
        # Log the error but continue to fallback methods
        import logging
        logging.getLogger(__name__).warning(f"PDF template overlay failed: {template_err}")
    
    # -------------------------------------------------------------------------
    # FALLBACK: AcroForm filling (if template_bytes is a fillable PDF)
    # -------------------------------------------------------------------------
    if isinstance(template_bytes, (bytes, bytearray)) and bytes(template_bytes)[:4] == b"%PDF":
        if progress_callback:
            try:
                progress_callback(55, "Trying AcroForm fill...")
            except Exception:
                pass
        filled = _fill_pdf_template_acroform(bytes(template_bytes), form_data)
        if filled is not None:
            return filled

    # -------------------------------------------------------------------------
    # FALLBACK: DOCX conversion (local MS Word/LibreOffice or CloudConvert)
    # -------------------------------------------------------------------------
    try:
        docx_bytes = generate_ptw_docx_from_template(template_bytes, form_data)
        if progress_callback:
            try:
                progress_callback(60, "Converting DOCX → PDF...")
            except Exception:
                pass
        pdf_bytes = convert_docx_to_pdf(docx_bytes, progress_callback=progress_callback)

        if isinstance(pdf_bytes, (bytes, bytearray)) and pdf_bytes[:4] == b"%PDF":
            return bytes(pdf_bytes)
    except Exception as docx_err:
        import logging
        logging.getLogger(__name__).warning(f"DOCX conversion fallback failed: {docx_err}")

    # -------------------------------------------------------------------------
    # All methods failed - provide actionable guidance
    # -------------------------------------------------------------------------
    raise RuntimeError(
        "PDF generation failed.\n\n"
        "Please ensure the PDF template is uploaded to Supabase Storage:\n"
        f"  Bucket: {SUPABASE_TEMPLATE_BUCKET}\n"
        f"  File: {TEMPLATE_PDF_FILE_NAME}\n\n"
        "The PDF template is required for hosted deployments."
    )


# =============================================================================
# FORM DATA BUILDER - Maps UI inputs to template placeholders
# =============================================================================


def build_form_data(
    # Header - linked to Work Order
    work_order_id: str,
    permit_validity_date: date,
    start_time: str,  # System-generated ISO timestamp
    end_time: str,    # System-generated ISO timestamp (start + 8 hours)
    site_name: str,   # Auto-filled from work_order
    work_location: str,  # Auto-derived: location + "-" + equipment
    work_description: str,
    contractor_name: str,
    # Hazards
    hz_live_dc_cables: bool,
    hz_loose_connectors: bool,
    hz_tracker_parts: bool,
    hz_dust: bool,
    hz_high_dc: bool,
    hz_poor_grounding: bool,
    hz_heavy_panels: bool,
    hz_wildlife: bool,
    hz_arc_flash: bool,
    hz_working_height: bool,
    hz_sharp_edges: bool,
    hz_lightning: bool,
    hz_improper_grounding: bool,
    hz_wet_surfaces: bool,
    hz_heat: bool,
    hz_overload: bool,
    hz_manual_handling: bool,
    hz_overhead_line: bool,
    hz_others_text: str,
    # Risks
    rk_electrocution: bool,
    rk_electric_shock: bool,
    rk_fall: bool,
    rk_tripping: bool,
    rk_burns: bool,
    rk_fire: bool,
    rk_back_injury: bool,
    rk_unexpected_energization: bool,
    rk_crushing: bool,
    rk_bites: bool,
    rk_heat_stress: bool,
    rk_electric_burn: bool,
    rk_falling_particles: bool,
    rk_others_text: str,
    # PPE
    ppe_helmet: bool,
    ppe_shoes: bool,
    ppe_reflective_vest: bool,
    ppe_goggles: bool,
    ppe_hrc_suit: bool,
    ppe_electrical_mat: bool,
    ppe_face_shield: bool,
    ppe_insulated_tools: bool,
    ppe_respirator: bool,
    ppe_dust_mask: bool,
    ppe_ear_plugs: bool,
    ppe_electrical_gloves: bool,
    ppe_harness: bool,
    ppe_lifeline: bool,
    ppe_cut_gloves: bool,
    ppe_others_text: str,
    # Safety Precautions
    sp_electrical_isolation: bool,
    sp_fire_extinguisher: bool,
    sp_proper_isolation: bool,
    sp_authorized_personnel: bool,
    sp_loto: bool,
    sp_signage: bool,
    sp_rescue_equipment: bool,
    sp_zero_voltage: bool,
    sp_pre_job_meeting: bool,
    sp_illumination: bool,
    sp_earthing: bool,
    sp_insulated_tools: bool,
    sp_escape_route: bool,
    sp_others_text: str,
    # Associated Permits
    ap_hot_work: bool,
    ap_general_work: bool,
    ap_loto: bool,
    ap_night_work: bool,
    ap_excavation: bool,
    ap_confined_space: bool,
    ap_height_work: bool,
    ap_lifting: bool,
    ap_others_text: str,
    # Tools/Equipment
    tools_equipment: str,
    # Issuer Checklist (Y/N/NA)
    chk_jsa: str,
    chk_environment: str,
    chk_loto: str,
    chk_firefighting: str,
    chk_energized_ppe: str,
    chk_rescue: str,
    chk_workers_fit: str,
    chk_grounded: str,
    chk_tools: str,
    chk_lighting: str,
    chk_rescue_plan: str,
    chk_signage: str,
    chk_testing_equipment: str,
    chk_conductive_removed: str,
    chk_line_clearance: str,
    chk_briefing: str,
    # Undertaking
    undertaking_accept: bool,
    # People (S1 level only: receiver + coworkers)
    receiver_name: str,
    coworker_1: str,
    coworker_2: str,
    coworker_3: str,
    coworker_4: str,
    coworker_5: str,
    coworker_6: str,
) -> dict:
    """
    Build the form_data dict for storage and document generation.
    
    S1 Level Only:
    - PTW is linked to a Work Order (work_order_id serves as permit reference)
    - Start/End times are system-controlled (captured at submit)
    - Site Name and Work Location are auto-derived from work order
    - Only Permit Receiver and Co-workers are filled at S1
    - Permit Holder, Issuer, Signatures, Closure are handled at S2 level
    
    Checkbox values are stored as booleans, converted to ticks for document.
    Checklist (Y/N/NA) values are stored and rendered as-is (bold in document).
    """
    # Raw form data for Supabase storage (preserves booleans)
    raw_data = {
        # Permit reference is the work_order_id
        "permit_no": work_order_id,
        "work_order_id": work_order_id,
        "permit_validity_date": str(permit_validity_date),
        "start_time": start_time,
        "end_time": end_time,
        "site_name": site_name,
        "work_location": work_location,
        "work_description": work_description,
        "contractor_name": contractor_name,
        # Hazards
        "hz_live_dc_cables": hz_live_dc_cables,
        "hz_loose_connectors": hz_loose_connectors,
        "hz_tracker_parts": hz_tracker_parts,
        "hz_dust": hz_dust,
        "hz_high_dc": hz_high_dc,
        "hz_poor_grounding": hz_poor_grounding,
        "hz_heavy_panels": hz_heavy_panels,
        "hz_wildlife": hz_wildlife,
        "hz_arc_flash": hz_arc_flash,
        "hz_working_height": hz_working_height,
        "hz_sharp_edges": hz_sharp_edges,
        "hz_lightning": hz_lightning,
        "hz_improper_grounding": hz_improper_grounding,
        "hz_wet_surfaces": hz_wet_surfaces,
        "hz_heat": hz_heat,
        "hz_overload": hz_overload,
        "hz_manual_handling": hz_manual_handling,
        "hz_overhead_line": hz_overhead_line,
        "hz_others_text": hz_others_text,
        # Risks
        "rk_electrocution": rk_electrocution,
        "rk_electric_shock": rk_electric_shock,
        "rk_fall": rk_fall,
        "rk_tripping": rk_tripping,
        "rk_burns": rk_burns,
        "rk_fire": rk_fire,
        "rk_back_injury": rk_back_injury,
        "rk_unexpected_energization": rk_unexpected_energization,
        "rk_crushing": rk_crushing,
        "rk_bites": rk_bites,
        "rk_heat_stress": rk_heat_stress,
        "rk_electric_burn": rk_electric_burn,
        "rk_falling_particles": rk_falling_particles,
        "rk_others_text": rk_others_text,
        # PPE
        "ppe_helmet": ppe_helmet,
        "ppe_shoes": ppe_shoes,
        "ppe_reflective_vest": ppe_reflective_vest,
        "ppe_goggles": ppe_goggles,
        "ppe_hrc_suit": ppe_hrc_suit,
        "ppe_electrical_mat": ppe_electrical_mat,
        "ppe_face_shield": ppe_face_shield,
        "ppe_insulated_tools": ppe_insulated_tools,
        "ppe_respirator": ppe_respirator,
        "ppe_dust_mask": ppe_dust_mask,
        "ppe_ear_plugs": ppe_ear_plugs,
        "ppe_electrical_gloves": ppe_electrical_gloves,
        "ppe_harness": ppe_harness,
        "ppe_lifeline": ppe_lifeline,
        "ppe_cut_gloves": ppe_cut_gloves,
        "ppe_others_text": ppe_others_text,
        # Safety Precautions
        "sp_electrical_isolation": sp_electrical_isolation,
        "sp_fire_extinguisher": sp_fire_extinguisher,
        "sp_proper_isolation": sp_proper_isolation,
        "sp_authorized_personnel": sp_authorized_personnel,
        "sp_loto": sp_loto,
        "sp_signage": sp_signage,
        "sp_rescue_equipment": sp_rescue_equipment,
        "sp_zero_voltage": sp_zero_voltage,
        "sp_pre_job_meeting": sp_pre_job_meeting,
        "sp_illumination": sp_illumination,
        "sp_earthing": sp_earthing,
        "sp_insulated_tools": sp_insulated_tools,
        "sp_escape_route": sp_escape_route,
        "sp_others_text": sp_others_text,
        # Associated Permits
        "ap_hot_work": ap_hot_work,
        "ap_general_work": ap_general_work,
        "ap_loto": ap_loto,
        "ap_night_work": ap_night_work,
        "ap_excavation": ap_excavation,
        "ap_confined_space": ap_confined_space,
        "ap_height_work": ap_height_work,
        "ap_lifting": ap_lifting,
        "ap_others_text": ap_others_text,
        # Tools
        "tools_equipment": tools_equipment,
        # Issuer Checklist (Y/N/NA stored as-is)
        "chk_jsa": chk_jsa,
        "chk_environment": chk_environment,
        "chk_loto": chk_loto,
        "chk_firefighting": chk_firefighting,
        "chk_energized_ppe": chk_energized_ppe,
        "chk_rescue": chk_rescue,
        "chk_workers_fit": chk_workers_fit,
        "chk_grounded": chk_grounded,
        "chk_tools": chk_tools,
        "chk_lighting": chk_lighting,
        "chk_rescue_plan": chk_rescue_plan,
        "chk_signage": chk_signage,
        "chk_testing_equipment": chk_testing_equipment,
        "chk_conductive_removed": chk_conductive_removed,
        "chk_line_clearance": chk_line_clearance,
        "chk_briefing": chk_briefing,
        # Undertaking
        "undertaking_accept": undertaking_accept,
        # People (S1 level only)
        "receiver_name": receiver_name,
        "coworker_1": coworker_1,
        "coworker_2": coworker_2,
        "coworker_3": coworker_3,
        "coworker_4": coworker_4,
        "coworker_5": coworker_5,
        "coworker_6": coworker_6,
        # S2 level fields - left blank at S1
        "holder_name": "",
        "issuer_name": "",
        # Metadata
        "submitted_at": datetime.now().isoformat(),
        # receiver_datetime must match start_time per PTW standards
        "receiver_datetime": start_time,
    }

    return raw_data


def build_doc_data(form_data: dict) -> dict:
    """
    Convert form_data to document-ready data.
    
    - Checkbox booleans are converted to tick marks
    - All values are stringified
    """
    doc_data = {}

    for k, v in form_data.items():
        # Checkbox fields (hz_, rk_, ppe_, sp_, ap_ prefixes + undertaking)
        if k.startswith(("hz_", "rk_", "ppe_", "sp_", "ap_")) and not k.endswith("_text"):
            doc_data[k] = _tick(v)
        elif k == "undertaking_accept":
            doc_data[k] = _tick(v)
        else:
            doc_data[k] = "" if v is None else str(v)

    return doc_data


# =============================================================================
# RENDER FUNCTION - Main UI
# =============================================================================


def render(db_path: str) -> None:
    """
    S1 Portal - Main render function.
    
    Tabs:
    - View Work Order: Filter and display work orders
    - Request PTW: Full Electrical PTW form
    - View Applied PTW: Browse and download submitted PTWs
    """
    # Hard access guard (role-based)
    from access_control import user_allowed_pages

    allowed_pages = user_allowed_pages()
    if "s1" not in allowed_pages:
        st.error("Access Denied - S1 Only")
        st.stop()

    st.markdown("# S1 Portal")

    # Import modern UI styles
    try:
        from modern_ui_styles import MODERN_UI_CSS
        st.markdown(MODERN_UI_CSS, unsafe_allow_html=True)
    except Exception:
        pass  # Fallback if import fails

    # Initialize session state for PTW
    if "s1_ptw_reset_counter" not in st.session_state:
        st.session_state["s1_ptw_reset_counter"] = 0
    if "s1_ptw_last_file" not in st.session_state:
        st.session_state["s1_ptw_last_file"] = None
    if "s1_ptw_last_permit_no" not in st.session_state:
        st.session_state["s1_ptw_last_permit_no"] = None
    if "s1_ptw_last_ext" not in st.session_state:
        st.session_state["s1_ptw_last_ext"] = "pdf"
    if "ptw_just_submitted" not in st.session_state:
        st.session_state["ptw_just_submitted"] = False

    # =========================================================================
    # Deterministic Navigation (no tab reset on rerun)
    # NOTE: Streamlit `st.tabs()` can jump back to the first tab after a button
    # click/form submit because the script reruns from top. We keep navigation
    # state explicitly in session_state and style it like tabs via CSS.
    # =========================================================================
    options = ["View Work Order", "Request PTW", "View Applied PTW", "Permit Closure"]
    if "s1_nav_page" not in st.session_state:
        st.session_state["s1_nav_page"] = options[0]

    # Prefer segmented control for a true "tab" look (and deterministic state).
    _seg = getattr(st, "segmented_control", None)
    if callable(_seg):
        try:
            active = _seg(" ", options=options, key="s1_nav_page", label_visibility="collapsed")
        except TypeError:
            # Older Streamlit: label_visibility may not exist
            active = _seg(" ", options=options, key="s1_nav_page")
    else:
        # Fallback: radio (still deterministic, but UI may vary by Streamlit version)
        active = st.radio(
            " ",
            options=options,
            horizontal=True,
            key="s1_nav_page",
            label_visibility="collapsed",
        )

    if active == "View Work Order":
        _render_view_work_order()
    elif active == "Request PTW":
        _render_request_ptw()
    elif active == "View Applied PTW":
        _render_view_applied_ptw()
    else:
        _render_permit_closure()


def _render_view_work_order() -> None:
    """Render the View Work Order tab."""
    st.markdown("## View Work Order")

    # Row hover styling
    st.markdown(
        """
        <style>
        .stDataFrame tbody tr:hover {
          background-color: rgba(148, 163, 184, 0.18) !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    sites = _list_sites_from_work_orders()
    if not sites:
        st.warning(
            "No Site Names could be loaded from `work_orders`.\n\n"
            "Common reasons:\n"
            "- `work_orders` is empty (ingestion not done)\n"
            "- Supabase RLS/policies are blocking SELECT access\n\n"
            "Fix: ensure your Supabase credentials have read access to `work_orders`."
        )
        return

    # Mount progress UI early (BEFORE any dependent filter queries) to avoid "page loads then progress".
    early_prog_slot = st.empty()
    early_msg_slot = st.empty()

    # Initialize session state
    if "s1_wo_last_df" not in st.session_state:
        st.session_state["s1_wo_last_df"] = None
        st.session_state["s1_wo_last_meta"] = None
    if "s1_wo_last_kpis" not in st.session_state:
        st.session_state["s1_wo_last_kpis"] = None
    if "s1_wo_run_fetch" not in st.session_state:
        st.session_state["s1_wo_run_fetch"] = False

    site_options = ["(select)"] + sites

    # Get current selections for dependent filters
    ss_site = st.session_state.get("s1_wo_site")
    ss_start = st.session_state.get("s1_wo_start")
    ss_end = st.session_state.get("s1_wo_end")

    have_site = ss_site not in (None, "(select)", "")
    have_dates = isinstance(ss_start, date) and isinstance(ss_end, date)

    # If we're about to fetch results (Submit clicked), skip dependent filter queries on this rerun
    # so the progress bar mounts immediately. (No query/business logic changes; just UI ordering.)
    if st.session_state.get("s1_wo_run_fetch"):
        locs, statuses_ui = [], []
    else:
        if have_site and have_dates:
            locs = _list_locations_from_work_orders(site_name=str(ss_site), start_date=ss_start, end_date=ss_end)
            statuses_ui = _list_statuses_from_work_orders(site_name=str(ss_site), start_date=ss_start, end_date=ss_end)
        else:
            locs = []
            statuses_ui = []

    loc_options = ["(all)"] + locs
    status_options = ["(all)"] + (statuses_ui if statuses_ui else UI_STATUSES)

    def _on_s1_wo_submit_click() -> None:
        # Flag allows next rerun to mount progress UI immediately before any heavy work.
        st.session_state["s1_wo_run_fetch"] = True

    # Filters OUTSIDE form to prevent Enter key triggering submit
    c1, c2, c3, c4, c5 = st.columns([2.0, 1.3, 1.3, 1.4, 1.4], vertical_alignment="bottom")
    with c1:
        site_name = st.selectbox("Site Name", options=site_options, index=0, key="s1_wo_site")
    with c2:
        start_date = st.date_input("Start Date", value=None, key="s1_wo_start")
    with c3:
        end_date = st.date_input("End Date", value=None, key="s1_wo_end")
    with c4:
        location = st.selectbox("Location", options=loc_options, index=0, key="s1_wo_location")
    with c5:
        status_ui = st.selectbox("Status", options=status_options, index=0, key="s1_wo_status")

    # Submit button (no form wrapper)
    submitted = st.button("Submit", type="primary", on_click=_on_s1_wo_submit_click)

    # If Submit was clicked, do the fetch with progress mounted at the top.
    if submitted or st.session_state.get("s1_wo_run_fetch"):
        # Ensure we only auto-run once per click
        st.session_state["s1_wo_run_fetch"] = False

        if not site_name or site_name == "(select)":
            early_prog_slot.empty()
            early_msg_slot.empty()
            st.error("Please select a Site Name.")
            return
        if start_date is None or end_date is None:
            early_prog_slot.empty()
            early_msg_slot.empty()
            st.error("Please select both Start Date and End Date.")
            return
        if start_date > end_date:
            early_prog_slot.empty()
            early_msg_slot.empty()
            st.error("Start Date must be on or before End Date.")
            return

        # Progress UX (mounted early)
        prog = early_prog_slot.progress(0, text="Safety First: Initializing...")
        early_msg_slot.caption("Safety First: Always verify permits and isolation before starting work.")
        _smooth_progress(prog, 0, 18, text="Validating filters...")
        _smooth_progress(prog, 18, 55, text="Fetching work orders...")

        loc_val = None if location in (None, "(all)") else location
        st_val = None if status_ui in (None, "(all)") else status_ui

        df = _fetch_work_orders(
            site_name=site_name,
            start_date=start_date,
            end_date=end_date,
            status_ui=st_val,
            location=loc_val,
        )

        _smooth_progress(prog, 55, 88, text="Preparing results...")
        _smooth_progress(prog, 88, 100, text="Results ready")
        early_prog_slot.empty()
        early_msg_slot.empty()

        # Persist results
        st.session_state["s1_wo_last_df"] = df
        st.session_state["s1_wo_last_meta"] = {
            "site_name": site_name,
            "start_date": start_date,
            "end_date": end_date,
            "status": st_val,
            "location": loc_val,
        }

        # Update KPIs
        if isinstance(df, pd.DataFrame) and not df.empty:
            total = int(len(df))
            c_rej = int((df["status"].astype("string").str.upper() == "REJECTED").sum())
            c_open = int((df["status"].astype("string").str.upper() == "OPEN").sum())
            c_wip = int((df["status"].astype("string").str.upper() == "WIP").sum())
            c_approved = int((df["status"].astype("string").str.upper() == "APPROVED").sum())
            c_closed = int((df["status"].astype("string").str.upper() == "CLOSED").sum())
        else:
            total = c_rej = c_open = c_wip = c_approved = c_closed = 0

        st.session_state["s1_wo_last_kpis"] = {
            "total": total,
            "rejected": c_rej,
            "open": c_open,
            "wip": c_wip,
            "approved": c_approved,
            "closed": c_closed,
        }

        if df.empty:
            st.info("No work orders found for the selected Site Name and Date Range.")
            return

    # Render cached results
    df_last = st.session_state.get("s1_wo_last_df")
    if isinstance(df_last, pd.DataFrame) and not df_last.empty:
        df = df_last

        # KPI cards
        k = st.session_state.get("s1_wo_last_kpis") or {}
        total = int(k.get("total", 0) or 0)
        c_rej = int(k.get("rejected", 0) or 0)
        c_open = int(k.get("open", 0) or 0)
        c_wip = int(k.get("wip", 0) or 0)
        c_approved = int(k.get("approved", 0) or 0)
        c_closed = int(k.get("closed", 0) or 0)

        st.markdown(
            """
            <style>
              .kpi-row { display: flex; gap: 12px; flex-wrap: wrap; margin: 10px 0 14px 0; }
              .kpi-card { flex: 1 1 160px; background: white; border: 1px solid rgba(148,163,184,0.35);
                          border-radius: 14px; padding: 14px 16px; box-shadow: 0 6px 18px rgba(15,23,42,0.06); }
              .kpi-title { font-size: 14px; color: #475569; margin-bottom: 6px; font-weight: 700; }
              .kpi-value { font-size: 34px; font-weight: 900; line-height: 1.05; }
            </style>
            """,
            unsafe_allow_html=True,
        )

        st.markdown(
            f"""
            <div class="kpi-row">
              <div class="kpi-card"><div class="kpi-title">Work Orders</div><div class="kpi-value" style="color:#2563eb;">{total}</div></div>
              <div class="kpi-card"><div class="kpi-title">Rejected</div><div class="kpi-value" style="color:#dc2626;">{c_rej}</div></div>
              <div class="kpi-card"><div class="kpi-title">Open</div><div class="kpi-value" style="color:#16a34a;">{c_open}</div></div>
              <div class="kpi-card"><div class="kpi-title">Awaiting Approval</div><div class="kpi-value" style="color:#f97316;">{c_wip}</div></div>
              <div class="kpi-card"><div class="kpi-title">Approved</div><div class="kpi-value" style="color:#10b981;">{c_approved}</div></div>
              <div class="kpi-card"><div class="kpi-title">Closed</div><div class="kpi-value" style="color:#065f46;">{c_closed}</div></div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Table with styling
        df_display = df.copy()
        df_display["status"] = df_display["status"].astype("string").fillna("").map(
            lambda s: DB_STATUS_TO_UI.get(str(s).strip().upper(), str(s))
        )
        styled = (
            df_display.style.map(_highlight_open_status, subset=["status"])
            .set_table_styles(
                [
                    {"selector": "th", "props": [("font-weight", "800"), ("color", "#0f172a")]},
                ]
            )
        )
        st.dataframe(styled, width="stretch", hide_index=True)


def _render_request_ptw() -> None:
    """
    Render the Request PTW tab with Work Order-driven flow.
    
    Flow:
    1. User selects a Date (to filter available work orders)
    2. User selects a Work Order ID from dropdown (only OPEN work orders for that date)
    3. Site Name and Work Location auto-fill from the selected work order
    4. Start/End times are system-controlled (captured at submit, end = start + 8 hours)
    5. At S1 level: only Permit Receiver and Co-workers are filled
    """
    # Enhanced UI Styling
    st.markdown("""
    <style>
        /* PTW Form Styling */
        .ptw-header {
            /* Light, modern header so text stays readable */
            background: linear-gradient(135deg, #e0f2fe 0%, #bfdbfe 100%);
            color: #0f172a;
            padding: 20px 25px;
            border-radius: 12px;
            margin-bottom: 20px;
            box-shadow: 0 6px 18px rgba(15, 23, 42, 0.10);
        }
        .ptw-header h2 {
            margin: 0;
            font-size: 24px;
            font-weight: 700;
        }
        .ptw-header p {
            margin: 8px 0 0 0;
            opacity: 0.9;
            font-size: 14px;
        }
        .section-card {
            background: #f8fafc;
            border: 1px solid #e2e8f0;
            border-radius: 10px;
            padding: 15px 20px;
            margin: 15px 0;
        }
        .section-title {
            color: #1e3a5f;
            font-weight: 700;
            font-size: 16px;
            margin-bottom: 10px;
            border-bottom: 2px solid #3b82f6;
            padding-bottom: 8px;
        }
        .success-card {
            background: linear-gradient(135deg, #059669 0%, #10b981 100%);
            color: white;
            padding: 25px;
            border-radius: 12px;
            text-align: center;
            margin: 20px 0;
        }
        .success-card h3 {
            margin: 0 0 10px 0;
            font-size: 22px;
        }
        .warning-banner {
            background: #fef3c7;
            border-left: 4px solid #f59e0b;
            padding: 12px 16px;
            border-radius: 0 8px 8px 0;
            margin: 10px 0;
        }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="ptw-header">
        <h2>⚡ Request PTW - Electrical Work Permit</h2>
        <p>Complete all sections below to submit a Permit To Work request</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Check if we just submitted successfully - show success message and reset prompt
    if st.session_state.get("ptw_just_submitted"):
        last_permit = st.session_state.get("s1_ptw_last_permit_no", "")
        
        # Success banner
        st.markdown(f"""
        <div class="success-card">
            <h3>✅ PTW Submitted Successfully!</h3>
            <p style="font-size: 18px; margin: 10px 0;"><strong>Permit No:</strong> {last_permit}</p>
            <p style="margin-top: 15px; font-size: 16px;">This PTW is under review with the supervisor.</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Action buttons
        col1, col2, col3 = st.columns([1, 2, 1])
        
        # Show download button FIRST (primary action)
        file_data = st.session_state.get("s1_ptw_last_file")
        if isinstance(file_data, (bytes, bytearray)) and len(file_data) > 0:
            with col2:
                st.download_button(
                    label="📥 Download PTW (PDF)",
                    data=file_data,
                    file_name=f"{last_permit}.pdf",
                    mime="application/pdf",
                    type="primary",
                    width="stretch",
                    key="download_submitted_ptw"
                )
        
        # Then show button to raise another PTW
        with col2:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("📝 Raise Another PTW Request", type="secondary", width="stretch", key="raise_another_ptw"):
                st.session_state["ptw_just_submitted"] = False
                _reset_ptw_form_state()
                st.rerun()
        
        return

    # Form key includes reset counter to force form reset
    reset_counter = st.session_state.get("s1_ptw_reset_counter", 0)
    form_key = f"ptw_electrical_form_{reset_counter}"

    # Initialize session state for work order selection (outside form for reactivity)
    if "ptw_selected_date" not in st.session_state:
        st.session_state["ptw_selected_date"] = date.today()
    if "ptw_selected_wo_id" not in st.session_state:
        st.session_state["ptw_selected_wo_id"] = None
    if "ptw_wo_details" not in st.session_state:
        st.session_state["ptw_wo_details"] = None
    if "ptw_wo_ids_for_date" not in st.session_state:
        st.session_state["ptw_wo_ids_for_date"] = []
    if "ptw_date_loading" not in st.session_state:
        st.session_state["ptw_date_loading"] = True  # mount progress on first load
    if "ptw_wo_loading" not in st.session_state:
        st.session_state["ptw_wo_loading"] = False

    # Use Streamlit fragments (if available) to prevent "whole page reload" feeling on
    # Date / Work Order selection. This is UI-only; backend logic stays identical.
    _frag = getattr(st, "fragment", None)

    def _render_work_order_selection_block() -> None:
        # ===== SECTION A: Permit Details (Work Order Selection) =====
        st.markdown('<div class="section-title">📌 A. Permit Details</div>', unsafe_allow_html=True)
        st.markdown(
            """
            <div style="background: #e0f2fe; border-left: 4px solid #0284c7; padding: 12px 16px; border-radius: 0 8px 8px 0; margin: 10px 0;">
                Select a date to view available Work Orders. The PTW will be linked to the selected Work Order.
            </div>
            """,
            unsafe_allow_html=True,
        )

        # Date selection (outside form for dynamic filtering)
        col_date, col_wo = st.columns([1, 2])
        with col_date:
            def _on_ptw_date_change() -> None:
                # UI-only: show progress immediately on rerun for date-driven work order list
                st.session_state["ptw_selected_date"] = st.session_state.get("ptw_date_selector")
                st.session_state["ptw_date_loading"] = True
                st.session_state["ptw_selected_wo_id"] = None
                st.session_state["ptw_wo_details"] = None
                st.session_state["ptw_site_name"] = ""
                st.session_state["ptw_work_location"] = ""

            selected_date = st.date_input(
                "Select Date (Work Order Date)",
                value=st.session_state.get("ptw_selected_date", date.today()),
                key="ptw_date_selector",
                on_change=_on_ptw_date_change,
            )

        # Mount loading UI BEFORE any backend work (prevents "page loads then progress")
        load_slot = st.empty()
        if st.session_state.get("ptw_date_loading"):
            prog = load_slot.progress(0, text="Loading work orders for selected date...")
            try:
                _smooth_progress(prog, 0, 35, text="Validating selection...")
                _smooth_progress(prog, 35, 80, text="Fetching OPEN work orders...")
                display_map = _build_work_order_display_map(selected_date)
                st.session_state["ptw_wo_display_map_for_date"] = display_map
                _smooth_progress(prog, 80, 100, text="Work orders ready")
            except Exception as e:
                st.error(f"Error fetching work orders: {e}")
                st.session_state["ptw_wo_ids_for_date"] = []
            finally:
                st.session_state["ptw_date_loading"] = False
                load_slot.empty()

        display_map = st.session_state.get("ptw_wo_display_map_for_date") or {}
        wo_labels = list(display_map.keys()) if isinstance(display_map, dict) else []

        with col_wo:
            if not wo_labels:
                st.warning(f"No OPEN work orders found for {selected_date.strftime('%Y-%m-%d')}")
                selected_ids: list[str] = []
            else:
                selected_labels = st.multiselect(
                    "Select Work Order(s)",
                    options=wo_labels,
                    key="ptw_wo_multiselect",
                )
                selected_ids = [str(display_map.get(lbl, "")).strip() for lbl in selected_labels if str(display_map.get(lbl, "")).strip()]

        # Fetch work order details for selected IDs (for auto-fill)
        wo_details_list: list[dict] = []
        if selected_ids:
            wo_prog_slot = st.empty()
            prog = wo_prog_slot.progress(0, text="Preparing selection...")
            try:
                _smooth_progress(prog, 0, 65, text="Fetching work order details...")
                for woid in selected_ids:
                    d = _get_work_order_details(woid)
                    if isinstance(d, dict) and d:
                        wo_details_list.append(d)
                _smooth_progress(prog, 65, 100, text="Ready")
            except Exception as e:
                st.error(f"Error fetching work order details: {e}")
                wo_details_list = []
            finally:
                wo_prog_slot.empty()

        site_name_display = ""
        work_location_display = ""
        permit_validity_display = ""
        
        if wo_details_list:
            site_names = sorted({str(d.get("site_name", "") or "").strip() for d in wo_details_list if str(d.get("site_name", "") or "").strip()})
            site_name_display = site_names[0] if site_names else ""

            # Auto-merge locations (multi-WO)
            locations = sorted({str(d.get("location", "") or "").strip() for d in wo_details_list if str(d.get("location", "") or "").strip()})
            if len(locations) == 1:
                work_location_display = locations[0]
            elif len(locations) > 1:
                work_location_display = " & ".join(locations)
            else:
                work_location_display = ""

            # Permit validity is CURRENT DATE (not work order date)
            permit_validity_display = date.today().strftime("%d-%m-%Y")

        # Persist auto-filled values for submit handler (UI-only ordering; same values)
        if selected_ids:
            st.session_state["ptw_selected_wo_ids"] = selected_ids
            # Backward compat: keep a primary ID (first) for legacy flows
            st.session_state["ptw_selected_wo_id"] = selected_ids[0]
            st.session_state["ptw_site_name"] = site_name_display
            st.session_state["ptw_work_location"] = work_location_display
            st.session_state["ptw_permit_no"] = "-".join(sorted(selected_ids))
        else:
            st.session_state["ptw_selected_wo_ids"] = []
            st.session_state["ptw_selected_wo_id"] = None
            st.session_state["ptw_site_name"] = ""
            st.session_state["ptw_work_location"] = ""
            st.session_state["ptw_permit_no"] = ""

        # Display auto-filled fields (read-only) - OUTSIDE form for reactive updates
        st.markdown("---")
        st.markdown("##### The following fields are auto-filled from the selected Work Order:")

        col_site, col_loc, col_validity = st.columns(3)
        with col_site:
            st.text_input(
                "Project / Site Name",
                value=site_name_display,
                disabled=True,
            )
        with col_loc:
            st.text_input(
                "Work Location",
                value=work_location_display,
                disabled=True,
            )
        with col_validity:
            st.text_input(
                "Permit Validity Date",
                value=permit_validity_display,
                disabled=True,
            )

        st.caption("Start Time and End Time are automatically recorded when you submit the PTW (End = Start + 8 hours).")

    # Wrap the WHOLE Request PTW interactive body in a fragment (if supported)
    # so Work Order selection triggers re-render of the dependent form below.
    def _render_request_ptw_interactive_body() -> None:
        _render_work_order_selection_block()

        # Read from session after selector block (fragment-safe)
        selected_ids = st.session_state.get("ptw_selected_wo_ids") or []
        selected_wo_id = st.session_state.get("ptw_selected_wo_id")
        
        # Permit validity is current date (NOT work order date)
        # As per requirement: validity = current datetime at submission
        permit_validity_date = date.today()

        # Check if we can proceed with the form
        if not selected_ids:
            st.warning("Please select one or more Work Orders to continue with the PTW request.")
            return

        # Store for use in form submission
        st.session_state["ptw_selected_wo_ids"] = selected_ids
        st.session_state["ptw_selected_wo_id"] = selected_wo_id
        st.session_state["ptw_site_name"] = st.session_state.get("ptw_site_name", "")
        st.session_state["ptw_work_location"] = st.session_state.get("ptw_work_location", "")

        # Prevent Enter key from implicitly submitting the form (UI-only safety)
        # Allow Enter inside TEXTAREA for multi-line notes.
        components.html(
            """
            <script>
              (function() {
                const handler = (e) => {
                  if (e.key === 'Enter') {
                    const tag = (e.target && e.target.tagName) ? e.target.tagName.toUpperCase() : '';
                    if (tag !== 'TEXTAREA') {
                      e.preventDefault();
                      e.stopPropagation();
                    }
                  }
                };
                window.removeEventListener('keydown', handler, true);
                window.addEventListener('keydown', handler, true);
              })();
            </script>
            """,
            height=0,
        )

        # If Submit PTW was clicked, avoid rebuilding the whole form before showing progress.
        if "ptw_submit_requested" not in st.session_state:
            st.session_state["ptw_submit_requested"] = False

        if st.session_state.get("ptw_submit_requested", False):
            st.session_state["ptw_submit_requested"] = False

            # Mount progress UI immediately (top of section)
            st.markdown('<div class="section-title">🚀 Submitting PTW</div>', unsafe_allow_html=True)

            def _ss(k: str, default: Any) -> Any:
                return st.session_state.get(k, default)

            work_order_ids = _ss("ptw_selected_wo_ids", []) or []
            permit_no = _ss("ptw_permit_no", "") or "-".join(sorted([str(x).strip() for x in work_order_ids if str(x).strip()]))
            site_name = _ss("ptw_site_name", "")
            work_location = _ss("ptw_work_location", "")
            
            # Permit validity is current date (NOT work order date)
            # As per requirement: validity = current datetime at submission
            permit_validity_date = date.today()

            _handle_ptw_submit(
                work_order_ids=work_order_ids,
                permit_no=permit_no,
                permit_validity_date=permit_validity_date,  # Current datetime
                site_name=site_name,
                work_location=work_location,
                work_description=_ss("ptw_work_desc", ""),
                contractor_name=_ss("ptw_contractor", ""),
                # Hazards
                hz_live_dc_cables=_ss("hz_live_dc_cables", False),
                hz_loose_connectors=_ss("hz_loose_connectors", False),
                hz_tracker_parts=_ss("hz_tracker_parts", False),
                hz_dust=_ss("hz_dust", False),
                hz_high_dc=_ss("hz_high_dc", False),
                hz_poor_grounding=_ss("hz_poor_grounding", False),
                hz_heavy_panels=_ss("hz_heavy_panels", False),
                hz_wildlife=_ss("hz_wildlife", False),
                hz_arc_flash=_ss("hz_arc_flash", False),
                hz_working_height=_ss("hz_working_height", False),
                hz_sharp_edges=_ss("hz_sharp_edges", False),
                hz_lightning=_ss("hz_lightning", False),
                hz_improper_grounding=_ss("hz_improper_grounding", False),
                hz_wet_surfaces=_ss("hz_wet_surfaces", False),
                hz_heat=_ss("hz_heat", False),
                hz_overload=_ss("hz_overload", False),
                hz_manual_handling=_ss("hz_manual_handling", False),
                hz_overhead_line=_ss("hz_overhead_line", False),
                hz_others_text=_ss("hz_others_text", ""),
                # Risks
                rk_electrocution=_ss("rk_electrocution", False),
                rk_electric_shock=_ss("rk_electric_shock", False),
                rk_fall=_ss("rk_fall", False),
                rk_tripping=_ss("rk_tripping", False),
                rk_burns=_ss("rk_burns", False),
                rk_fire=_ss("rk_fire", False),
                rk_back_injury=_ss("rk_back_injury", False),
                rk_unexpected_energization=_ss("rk_unexpected_energization", False),
                rk_crushing=_ss("rk_crushing", False),
                rk_bites=_ss("rk_bites", False),
                rk_heat_stress=_ss("rk_heat_stress", False),
                rk_electric_burn=_ss("rk_electric_burn", False),
                rk_falling_particles=_ss("rk_falling_particles", False),
                rk_others_text=_ss("rk_others_text", ""),
                # PPE
                ppe_helmet=_ss("ppe_helmet", False),
                ppe_shoes=_ss("ppe_shoes", False),
                ppe_reflective_vest=_ss("ppe_reflective_vest", False),
                ppe_goggles=_ss("ppe_goggles", False),
                ppe_hrc_suit=_ss("ppe_hrc_suit", False),
                ppe_electrical_mat=_ss("ppe_electrical_mat", False),
                ppe_face_shield=_ss("ppe_face_shield", False),
                ppe_insulated_tools=_ss("ppe_insulated_tools", False),
                ppe_respirator=_ss("ppe_respirator", False),
                ppe_dust_mask=_ss("ppe_dust_mask", False),
                ppe_ear_plugs=_ss("ppe_ear_plugs", False),
                ppe_electrical_gloves=_ss("ppe_electrical_gloves", False),
                ppe_harness=_ss("ppe_harness", False),
                ppe_lifeline=_ss("ppe_lifeline", False),
                ppe_cut_gloves=_ss("ppe_cut_gloves", False),
                ppe_others_text=_ss("ppe_others_text", ""),
                # Safety Precautions
                sp_electrical_isolation=_ss("sp_electrical_isolation", False),
                sp_fire_extinguisher=_ss("sp_fire_extinguisher", False),
                sp_proper_isolation=_ss("sp_proper_isolation", False),
                sp_authorized_personnel=_ss("sp_authorized_personnel", False),
                sp_loto=_ss("sp_loto", False),
                sp_signage=_ss("sp_signage", False),
                sp_rescue_equipment=_ss("sp_rescue_equipment", False),
                sp_zero_voltage=_ss("sp_zero_voltage", False),
                sp_pre_job_meeting=_ss("sp_pre_job_meeting", False),
                sp_illumination=_ss("sp_illumination", False),
                sp_earthing=_ss("sp_earthing", False),
                sp_insulated_tools=_ss("sp_insulated_tools", False),
                sp_escape_route=_ss("sp_escape_route", False),
                sp_others_text=_ss("sp_others_text", ""),
                # Associated Permits
                ap_hot_work=_ss("ap_hot_work", False),
                ap_general_work=_ss("ap_general_work", False),
                ap_loto=_ss("ap_loto", False),
                ap_night_work=_ss("ap_night_work", False),
                ap_excavation=_ss("ap_excavation", False),
                ap_confined_space=_ss("ap_confined_space", False),
                ap_height_work=_ss("ap_height_work", False),
                ap_lifting=_ss("ap_lifting", False),
                ap_others_text=_ss("ap_others_text", ""),
                # Tools (Section G uses key="ptw_tools_equipment")
                tools_equipment=_ss("ptw_tools_equipment", ""),
                # Checklist (Y/N/NA)
                chk_jsa=_ss("chk_jsa", None),
                chk_environment=_ss("chk_environment", None),
                chk_loto=_ss("chk_loto", None),
                chk_firefighting=_ss("chk_firefighting", None),
                chk_energized_ppe=_ss("chk_energized_ppe", None),
                chk_rescue=_ss("chk_rescue", None),
                chk_workers_fit=_ss("chk_workers_fit", None),
                chk_grounded=_ss("chk_grounded", None),
                chk_tools=_ss("chk_tools", None),
                chk_lighting=_ss("chk_lighting", None),
                chk_rescue_plan=_ss("chk_rescue_plan", None),
                chk_signage=_ss("chk_signage", None),
                chk_testing_equipment=_ss("chk_testing_equipment", None),
                chk_conductive_removed=_ss("chk_conductive_removed", None),
                chk_line_clearance=_ss("chk_line_clearance", None),
                chk_briefing=_ss("chk_briefing", None),
                # Undertaking
                undertaking_accept=_ss("ptw_undertaking", False),
                # People
                receiver_name=_ss("ptw_receiver", ""),
                coworker_1=_ss("ptw_coworker1", ""),
                coworker_2=_ss("ptw_coworker2", ""),
                coworker_3=_ss("ptw_coworker3", ""),
                coworker_4=_ss("ptw_coworker4", ""),
                coworker_5=_ss("ptw_coworker5", ""),
                coworker_6=_ss("ptw_coworker6", ""),
            )
            st.stop()

        # Otherwise, continue with normal form rendering below (existing code)

    if callable(_frag):
        st.fragment(_render_request_ptw_interactive_body)()
    else:
        _render_request_ptw_interactive_body()

    with st.form(form_key):
        # Manual entry fields with better styling
        st.markdown('<div class="section-title">📋 Work Details</div>', unsafe_allow_html=True)
        work_description = st.text_area("Description of Work *", height=100, key="ptw_work_desc", 
                                         placeholder="Describe the work to be performed...")
        contractor_name = st.text_area("Contractor Name", key="ptw_contractor",
                                        placeholder="Enter contractor/company name", height=38)

        st.divider()

        # ===== SECTION B: Hazards =====
        st.markdown('<div class="section-title">⚠️ B. Hazards / Hazardous Activities</div>', unsafe_allow_html=True)
        st.caption("Select all applicable hazards identified for this work")

        hz_cols = st.columns(4)
        with hz_cols[0]:
            hz_live_dc_cables = st.checkbox("Live DC cables", key="hz_live_dc_cables")
            hz_high_dc = st.checkbox("High DC voltage", key="hz_high_dc")
            hz_arc_flash = st.checkbox("Arc flash / short circuit", key="hz_arc_flash")
            hz_improper_grounding = st.checkbox("Improper grounding", key="hz_improper_grounding")
            hz_overload = st.checkbox("Overload", key="hz_overload")
        with hz_cols[1]:
            hz_loose_connectors = st.checkbox("Loose connectors", key="hz_loose_connectors")
            hz_poor_grounding = st.checkbox("Poor grounding", key="hz_poor_grounding")
            hz_working_height = st.checkbox("Working at height", key="hz_working_height")
            hz_wet_surfaces = st.checkbox("Wet surfaces", key="hz_wet_surfaces")
            hz_manual_handling = st.checkbox("Manual handling", key="hz_manual_handling")
        with hz_cols[2]:
            hz_tracker_parts = st.checkbox("Tracker moving parts", key="hz_tracker_parts")
            hz_heavy_panels = st.checkbox("Heavy panels", key="hz_heavy_panels")
            hz_sharp_edges = st.checkbox("Sharp edges", key="hz_sharp_edges")
            hz_heat = st.checkbox("Heat", key="hz_heat")
            hz_overhead_line = st.checkbox("Overhead line", key="hz_overhead_line")
        with hz_cols[3]:
            hz_dust = st.checkbox("Dust", key="hz_dust")
            hz_wildlife = st.checkbox("Wildlife", key="hz_wildlife")
            hz_lightning = st.checkbox("Lightning", key="hz_lightning")

        hz_others_text = st.text_area("Other Hazards (if any)", key="hz_others_text", height=38)

        st.divider()

        # ===== SECTION C: Risk Identification =====
        st.markdown('<div class="section-title">🎯 C. Risk Identification</div>', unsafe_allow_html=True)
        st.caption("Select all identified risks for this work")

        rk_cols = st.columns(4)
        with rk_cols[0]:
            rk_electrocution = st.checkbox("Electrocution", key="rk_electrocution")
            rk_electric_shock = st.checkbox("Electric shock", key="rk_electric_shock")
            rk_electric_burn = st.checkbox("Electric burn", key="rk_electric_burn")
            rk_unexpected_energization = st.checkbox("Unexpected energization", key="rk_unexpected_energization")
        with rk_cols[1]:
            rk_fall = st.checkbox("Fall", key="rk_fall")
            rk_tripping = st.checkbox("Tripping", key="rk_tripping")
            rk_falling_particles = st.checkbox("Falling particles", key="rk_falling_particles")
            rk_crushing = st.checkbox("Crushing", key="rk_crushing")
        with rk_cols[2]:
            rk_burns = st.checkbox("Burns", key="rk_burns")
            rk_fire = st.checkbox("Fire", key="rk_fire")
            rk_heat_stress = st.checkbox("Heat stress", key="rk_heat_stress")
            rk_bites = st.checkbox("Bites/Stings", key="rk_bites")
        with rk_cols[3]:
            rk_back_injury = st.checkbox("Back injury", key="rk_back_injury")

        rk_others_text = st.text_area("Other Risks (if any)", key="rk_others_text", height=38)

        st.divider()

        # ===== SECTION D: PPE =====
        st.markdown('<div class="section-title">🦺 D. Personal Protective Equipment (PPE)</div>', unsafe_allow_html=True)
        st.caption("Select all required PPE for this work")

        ppe_cols = st.columns(4)
        with ppe_cols[0]:
            ppe_helmet = st.checkbox("Safety Helmet", key="ppe_helmet")
            ppe_shoes = st.checkbox("Safety Shoes", key="ppe_shoes")
            ppe_electrical_gloves = st.checkbox("Electrical Gloves", key="ppe_electrical_gloves")
            ppe_harness = st.checkbox("Safety Harness", key="ppe_harness")
        with ppe_cols[1]:
            ppe_reflective_vest = st.checkbox("Reflective Vest", key="ppe_reflective_vest")
            ppe_goggles = st.checkbox("Safety Goggles", key="ppe_goggles")
            ppe_face_shield = st.checkbox("Face Shield", key="ppe_face_shield")
            ppe_lifeline = st.checkbox("Lifeline", key="ppe_lifeline")
        with ppe_cols[2]:
            ppe_hrc_suit = st.checkbox("HRC Suit", key="ppe_hrc_suit")
            ppe_electrical_mat = st.checkbox("Electrical Mat", key="ppe_electrical_mat")
            ppe_insulated_tools = st.checkbox("Insulated Tools", key="ppe_insulated_tools")
            ppe_cut_gloves = st.checkbox("Cut Resistant Gloves", key="ppe_cut_gloves")
        with ppe_cols[3]:
            ppe_respirator = st.checkbox("Respirator", key="ppe_respirator")
            ppe_dust_mask = st.checkbox("Dust Mask", key="ppe_dust_mask")
            ppe_ear_plugs = st.checkbox("Ear Plugs", key="ppe_ear_plugs")

        ppe_others_text = st.text_area("Other PPE (if any)", key="ppe_others_text", height=38)

        st.divider()

        # ===== SECTION E: Safety Precautions =====
        st.markdown('<div class="section-title">🛡️ E. Safety Precautions</div>', unsafe_allow_html=True)
        st.caption("Select all applicable safety precautions")

        sp_cols = st.columns(4)
        with sp_cols[0]:
            sp_electrical_isolation = st.checkbox("Electrical isolation", key="sp_electrical_isolation")
            sp_proper_isolation = st.checkbox("Proper isolation", key="sp_proper_isolation")
            sp_loto = st.checkbox("LOTO applied", key="sp_loto")
            sp_earthing = st.checkbox("Earthing", key="sp_earthing")
        with sp_cols[1]:
            sp_fire_extinguisher = st.checkbox("Fire extinguisher", key="sp_fire_extinguisher")
            sp_authorized_personnel = st.checkbox("Authorized personnel only", key="sp_authorized_personnel")
            sp_signage = st.checkbox("Signage placed", key="sp_signage")
            sp_insulated_tools = st.checkbox("Insulated tools", key="sp_insulated_tools")
        with sp_cols[2]:
            sp_rescue_equipment = st.checkbox("Rescue equipment", key="sp_rescue_equipment")
            sp_zero_voltage = st.checkbox("Zero voltage verified", key="sp_zero_voltage")
            sp_pre_job_meeting = st.checkbox("Pre-job meeting", key="sp_pre_job_meeting")
            sp_escape_route = st.checkbox("Escape route clear", key="sp_escape_route")
        with sp_cols[3]:
            sp_illumination = st.checkbox("Adequate illumination", key="sp_illumination")

        sp_others_text = st.text_area("Other Precautions (if any)", key="sp_others_text", height=38)

        st.divider()

        # ===== SECTION F: Associated Permits =====
        st.markdown('<div class="section-title">📄 F. Associated Permits</div>', unsafe_allow_html=True)
        st.caption("Select any associated permits required")

        ap_cols = st.columns(4)
        with ap_cols[0]:
            ap_hot_work = st.checkbox("Hot Work Permit", key="ap_hot_work")
            ap_loto = st.checkbox("LOTO Permit", key="ap_loto")
            ap_height_work = st.checkbox("Height Work Permit", key="ap_height_work")
        with ap_cols[1]:
            ap_general_work = st.checkbox("General Work Permit", key="ap_general_work")
            ap_night_work = st.checkbox("Night Work Permit", key="ap_night_work")
            ap_lifting = st.checkbox("Lifting Permit", key="ap_lifting")
        with ap_cols[2]:
            ap_excavation = st.checkbox("Excavation Permit", key="ap_excavation")
            ap_confined_space = st.checkbox("Confined Space Permit", key="ap_confined_space")

        ap_others_text = st.text_area("Other Permits (if any)", key="ap_others_text", height=38)

        st.divider()

        # ===== SECTION G: Tools/Equipment =====
        st.markdown('<div class="section-title">🔧 G. Tools / Equipment Required</div>', unsafe_allow_html=True)
        tools_equipment = st.text_area(
            "List all tools and equipment required for the work",
            height=80,
            key="ptw_tools_equipment"
        )

        st.divider()

        # ===== SECTION H: Issuer Safety Checklist =====
        st.markdown('<div class="section-title">✅ H. Safety Checklist (Permit Issuer)</div>', unsafe_allow_html=True)
        st.caption("Complete all safety checks before issuing permit")

        chk_col1, chk_col2 = st.columns(2)
        with chk_col1:
            chk_jsa = st.radio("Is JSA carried out?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_jsa")
            chk_environment = st.radio("Is environment suitable?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_environment")
            chk_loto = st.radio("Is LOTO applied?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_loto")
            chk_firefighting = st.radio("Firefighting equipment available?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_firefighting")
            chk_energized_ppe = st.radio("PPE for energized work?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_energized_ppe")
            chk_rescue = st.radio("Rescue equipment available?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_rescue")
            chk_workers_fit = st.radio("Workers medically fit?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_workers_fit")
            chk_grounded = st.radio("Equipment grounded?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_grounded")

        with chk_col2:
            chk_tools = st.radio("Tools inspected?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_tools")
            chk_lighting = st.radio("Adequate lighting?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_lighting")
            chk_rescue_plan = st.radio("Rescue plan in place?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_rescue_plan")
            chk_signage = st.radio("Warning signage placed?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_signage")
            chk_testing_equipment = st.radio("Testing equipment calibrated?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_testing_equipment")
            chk_conductive_removed = st.radio("Conductive items removed?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_conductive_removed")
            chk_line_clearance = st.radio("Line clearance obtained?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_line_clearance")
            chk_briefing = st.radio("Safety briefing done?", ["Y", "N", "NA"], index=None, horizontal=True, key="chk_briefing")

        st.divider()

        # ===== SECTION I: Undertaking =====
        st.markdown('<div class="section-title">📝 I. Undertaking (MANDATORY)</div>', unsafe_allow_html=True)
        st.markdown("""
        <div class="warning-banner">
            <strong>⚠️ Important:</strong> This undertaking must be accepted before submitting the permit.
        </div>
        """, unsafe_allow_html=True)
        undertaking_accept = st.checkbox(
            "I have reviewed and understood the risk assessment, safety precautions, and emergency procedures. "
            "I accept responsibility for ensuring all personnel involved comply with the permit conditions.",
            key="ptw_undertaking"
        )

        st.divider()

        # ===== SECTION J: People Involved (S1 Level) =====
        st.markdown('<div class="section-title">👥 J. People Involved</div>', unsafe_allow_html=True)
        st.caption("At S1 (Request) level: Enter Permit Receiver and Co-workers. Permit Holder and Issuer are assigned at S2 (Approval) level.")

        ppl_col1, ppl_col2, ppl_col3 = st.columns(3)
        with ppl_col1:
            receiver_name = st.text_area("Permit Receiver Name *", key="ptw_receiver", height=38)
            coworker_1 = st.text_area("Co-worker 1", key="ptw_coworker1", height=38)
        with ppl_col2:
            coworker_2 = st.text_area("Co-worker 2", key="ptw_coworker2", height=38)
            coworker_3 = st.text_area("Co-worker 3", key="ptw_coworker3", height=38)
        with ppl_col3:
            coworker_4 = st.text_area("Co-worker 4", key="ptw_coworker4", height=38)
            coworker_5 = st.text_area("Co-worker 5", key="ptw_coworker5", height=38)

        coworker_6 = st.text_area("Co-worker 6", key="ptw_coworker6", height=38)

        st.divider()

        # ===== SUBMIT BUTTON =====
        is_submitting = bool(st.session_state.get("ptw_is_submitting", False))
        def _on_ptw_submit_click() -> None:
            # UI-only: next rerun will show progress immediately and run submit without rebuilding the full form first.
            st.session_state["ptw_submit_requested"] = True
        submitted = st.form_submit_button(
            "Submit PTW",
            type="primary",
            width="stretch",
            disabled=is_submitting,
            on_click=_on_ptw_submit_click,
        )

    # Submission is handled via ptw_submit_requested flag (above) to ensure progress appears immediately.

    # Download button is now shown in the success screen (ptw_just_submitted state)


def _handle_ptw_submit(**kwargs) -> None:
    """
    Handle PTW form submission with progress feedback.
    
    Key behaviors:
    - work_order_id serves as permit_no (PTW linked to Work Order)
    - Start Time is captured at moment of submit
    - End Time = Start Time + 8 hours (system-controlled)
    - PDF-only output (no DOCX download exposed)
    """

    # Validation
    if not kwargs.get("undertaking_accept"):
        st.error("Undertaking acceptance is mandatory. Please accept the undertaking to proceed.")
        return

    work_order_ids = kwargs.get("work_order_ids") or []
    if not isinstance(work_order_ids, list):
        work_order_ids = []
    work_order_ids = [str(x).strip() for x in work_order_ids if str(x).strip()]
    if not work_order_ids:
        st.error("Work Order selection is required. Please select one or more Work Orders.")
        return

    if not kwargs.get("site_name"):
        st.error("Site Name could not be determined. Please select a valid Work Order.")
        return

    if not kwargs.get("receiver_name"):
        st.error("Permit Receiver Name is required.")
        return

    # Progress/status placeholders (keeps layout stable)
    progress_placeholder = st.empty()
    status_placeholder = st.empty()
    msg_placeholder = st.empty()

    try:
        # Prevent double clicks / duplicate submits in same session
        if st.session_state.get("ptw_is_submitting", False):
            return
        st.session_state["ptw_is_submitting"] = True

        progress_placeholder.progress(0, text="Validating inputs...")
        status_placeholder.info("Step 1/4: Validating form data")

        # System-controlled timestamps (IST timezone)
        submit_time = datetime.now(ZoneInfo("Asia/Kolkata"))
        start_time_str = submit_time.strftime("%H:%M:%S")
        end_time = submit_time + timedelta(hours=8)
        end_time_str = end_time.strftime("%H:%M:%S")

        permit_no = str(kwargs.get("permit_no") or "").strip() or "-".join(sorted(work_order_ids))

        # Block duplicate PTW for any selected work_order_id (safety requirement)
        exists, existing_ptw_id = _ptw_exists_for_work_orders(work_order_ids)
        if exists:
            progress_placeholder.empty()
            status_placeholder.empty()
            msg_placeholder.error(
                "One or more selected Work Orders already have a PTW."
                + (f" Existing PTW ID: `{existing_ptw_id}`" if existing_ptw_id else "")
            )
            st.session_state["ptw_is_submitting"] = False
            return

        # Build form data with system-controlled times
        form_data = build_form_data(
            work_order_id=permit_no,
            permit_validity_date=kwargs["permit_validity_date"],
            start_time=start_time_str,
            end_time=end_time_str,
            site_name=kwargs["site_name"],
            work_location=kwargs["work_location"],
            work_description=kwargs.get("work_description", ""),
            contractor_name=kwargs.get("contractor_name", ""),
            # Hazards
            hz_live_dc_cables=kwargs.get("hz_live_dc_cables", False),
            hz_loose_connectors=kwargs.get("hz_loose_connectors", False),
            hz_tracker_parts=kwargs.get("hz_tracker_parts", False),
            hz_dust=kwargs.get("hz_dust", False),
            hz_high_dc=kwargs.get("hz_high_dc", False),
            hz_poor_grounding=kwargs.get("hz_poor_grounding", False),
            hz_heavy_panels=kwargs.get("hz_heavy_panels", False),
            hz_wildlife=kwargs.get("hz_wildlife", False),
            hz_arc_flash=kwargs.get("hz_arc_flash", False),
            hz_working_height=kwargs.get("hz_working_height", False),
            hz_sharp_edges=kwargs.get("hz_sharp_edges", False),
            hz_lightning=kwargs.get("hz_lightning", False),
            hz_improper_grounding=kwargs.get("hz_improper_grounding", False),
            hz_wet_surfaces=kwargs.get("hz_wet_surfaces", False),
            hz_heat=kwargs.get("hz_heat", False),
            hz_overload=kwargs.get("hz_overload", False),
            hz_manual_handling=kwargs.get("hz_manual_handling", False),
            hz_overhead_line=kwargs.get("hz_overhead_line", False),
            hz_others_text=kwargs.get("hz_others_text", ""),
            # Risks
            rk_electrocution=kwargs.get("rk_electrocution", False),
            rk_electric_shock=kwargs.get("rk_electric_shock", False),
            rk_fall=kwargs.get("rk_fall", False),
            rk_tripping=kwargs.get("rk_tripping", False),
            rk_burns=kwargs.get("rk_burns", False),
            rk_fire=kwargs.get("rk_fire", False),
            rk_back_injury=kwargs.get("rk_back_injury", False),
            rk_unexpected_energization=kwargs.get("rk_unexpected_energization", False),
            rk_crushing=kwargs.get("rk_crushing", False),
            rk_bites=kwargs.get("rk_bites", False),
            rk_heat_stress=kwargs.get("rk_heat_stress", False),
            rk_electric_burn=kwargs.get("rk_electric_burn", False),
            rk_falling_particles=kwargs.get("rk_falling_particles", False),
            rk_others_text=kwargs.get("rk_others_text", ""),
            # PPE
            ppe_helmet=kwargs.get("ppe_helmet", False),
            ppe_shoes=kwargs.get("ppe_shoes", False),
            ppe_reflective_vest=kwargs.get("ppe_reflective_vest", False),
            ppe_goggles=kwargs.get("ppe_goggles", False),
            ppe_hrc_suit=kwargs.get("ppe_hrc_suit", False),
            ppe_electrical_mat=kwargs.get("ppe_electrical_mat", False),
            ppe_face_shield=kwargs.get("ppe_face_shield", False),
            ppe_insulated_tools=kwargs.get("ppe_insulated_tools", False),
            ppe_respirator=kwargs.get("ppe_respirator", False),
            ppe_dust_mask=kwargs.get("ppe_dust_mask", False),
            ppe_ear_plugs=kwargs.get("ppe_ear_plugs", False),
            ppe_electrical_gloves=kwargs.get("ppe_electrical_gloves", False),
            ppe_harness=kwargs.get("ppe_harness", False),
            ppe_lifeline=kwargs.get("ppe_lifeline", False),
            ppe_cut_gloves=kwargs.get("ppe_cut_gloves", False),
            ppe_others_text=kwargs.get("ppe_others_text", ""),
            # Safety Precautions
            sp_electrical_isolation=kwargs.get("sp_electrical_isolation", False),
            sp_fire_extinguisher=kwargs.get("sp_fire_extinguisher", False),
            sp_proper_isolation=kwargs.get("sp_proper_isolation", False),
            sp_authorized_personnel=kwargs.get("sp_authorized_personnel", False),
            sp_loto=kwargs.get("sp_loto", False),
            sp_signage=kwargs.get("sp_signage", False),
            sp_rescue_equipment=kwargs.get("sp_rescue_equipment", False),
            sp_zero_voltage=kwargs.get("sp_zero_voltage", False),
            sp_pre_job_meeting=kwargs.get("sp_pre_job_meeting", False),
            sp_illumination=kwargs.get("sp_illumination", False),
            sp_earthing=kwargs.get("sp_earthing", False),
            sp_insulated_tools=kwargs.get("sp_insulated_tools", False),
            sp_escape_route=kwargs.get("sp_escape_route", False),
            sp_others_text=kwargs.get("sp_others_text", ""),
            # Associated Permits
            ap_hot_work=kwargs.get("ap_hot_work", False),
            ap_general_work=kwargs.get("ap_general_work", False),
            ap_loto=kwargs.get("ap_loto", False),
            ap_night_work=kwargs.get("ap_night_work", False),
            ap_excavation=kwargs.get("ap_excavation", False),
            ap_confined_space=kwargs.get("ap_confined_space", False),
            ap_height_work=kwargs.get("ap_height_work", False),
            ap_lifting=kwargs.get("ap_lifting", False),
            ap_others_text=kwargs.get("ap_others_text", ""),
            # Tools
            tools_equipment=kwargs.get("tools_equipment", ""),
            # Checklist (Y/N/NA)
            chk_jsa=kwargs.get("chk_jsa") or "NA",
            chk_environment=kwargs.get("chk_environment") or "NA",
            chk_loto=kwargs.get("chk_loto") or "NA",
            chk_firefighting=kwargs.get("chk_firefighting") or "NA",
            chk_energized_ppe=kwargs.get("chk_energized_ppe") or "NA",
            chk_rescue=kwargs.get("chk_rescue") or "NA",
            chk_workers_fit=kwargs.get("chk_workers_fit") or "NA",
            chk_grounded=kwargs.get("chk_grounded") or "NA",
            chk_tools=kwargs.get("chk_tools") or "NA",
            chk_lighting=kwargs.get("chk_lighting") or "NA",
            chk_rescue_plan=kwargs.get("chk_rescue_plan") or "NA",
            chk_signage=kwargs.get("chk_signage") or "NA",
            chk_testing_equipment=kwargs.get("chk_testing_equipment") or "NA",
            chk_conductive_removed=kwargs.get("chk_conductive_removed") or "NA",
            chk_line_clearance=kwargs.get("chk_line_clearance") or "NA",
            chk_briefing=kwargs.get("chk_briefing") or "NA",
            # Undertaking
            undertaking_accept=kwargs.get("undertaking_accept", False),
            # People (S1 level only)
            receiver_name=kwargs.get("receiver_name", ""),
            coworker_1=kwargs.get("coworker_1", ""),
            coworker_2=kwargs.get("coworker_2", ""),
            coworker_3=kwargs.get("coworker_3", ""),
            coworker_4=kwargs.get("coworker_4", ""),
            coworker_5=kwargs.get("coworker_5", ""),
            coworker_6=kwargs.get("coworker_6", ""),
        )
        # Store multi-WO coverage (read-only for other portals; no schema changes)
        form_data["work_order_ids"] = work_order_ids

        # Signature auto-sync (do not ask separately)
        recv_name = str(form_data.get("receiver_name") or "").strip()
        if recv_name:
            form_data["receiver_signature"] = recv_name
            form_data["permit_receiver_signature"] = recv_name
        # Tools overlay helper expects this key (keeps existing tools_equipment intact)
        form_data["tools_equipment_required"] = form_data.get("tools_equipment", "")
        # Undertaking overlay helper expects this key
        form_data["undertaking_checkbox"] = bool(form_data.get("undertaking_accept", False))

        progress_placeholder.progress(25, text="Saving to Supabase...")
        status_placeholder.info("Step 2/4: Saving PTW to database")

        # Insert into Supabase (permit_no = combined work_order_ids)
        ptw_id = insert_ptw_request(
            permit_no=permit_no,
            site_name=kwargs["site_name"],
            created_by=kwargs["receiver_name"],
            form_data=form_data,
        )
        
        # Update work_orders.date_s1_created to mark PTW submission time
        # This drives the status derivation (OPEN -> WIP)
        for wo in work_order_ids:
            _update_work_order_s1_created(
                work_order_id=wo,
                s1_timestamp=submit_time.isoformat(sep=" ", timespec="seconds"),
            )

        progress_placeholder.progress(50, text="Downloading template from Supabase...")
        status_placeholder.info("Step 3/4: Generating PTW document")

        # Download template from Supabase storage
        try:
            template_bytes = _download_template_from_supabase()
        except Exception as e:
            raise RuntimeError(f"Failed to download template: {e}") from e

        # Generate document (with attachments if S2's generator is available)
        try:
            from S2 import generate_ptw_pdf_with_attachments
            # Note: at S1 submission, no evidence files exist yet, but using this generator
            # ensures consistency. Attachments page will be empty for S1-only PDFs.
            pdf_bytes = generate_ptw_pdf_with_attachments(
                form_data=form_data,
                work_order_id=permit_no,
                progress_callback=None,
            )
        except Exception:
            # Fallback: basic PDF without attachments
            doc_data = build_doc_data(form_data)
            pdf_bytes = generate_ptw_pdf(template_bytes, doc_data)  # PDF-only (raises if fails)
            pdf_bytes = _add_overlay_elements(pdf_bytes, form_data)

        progress_placeholder.progress(75, text="Finalizing...")
        status_placeholder.info("Step 4/4: Finalizing submission")

        # Store for download (PDF only)
        st.session_state["s1_ptw_last_file"] = pdf_bytes
        st.session_state["s1_ptw_last_permit_no"] = permit_no
        st.session_state["s1_ptw_last_ext"] = "pdf"

        # Reset form for next entry
        st.session_state["s1_ptw_reset_counter"] = st.session_state.get("s1_ptw_reset_counter", 0) + 1
        # Clear work order selection
        st.session_state["ptw_selected_wo_id"] = None
        st.session_state["ptw_wo_details"] = None

        progress_placeholder.progress(100, text="Complete!")
        status_placeholder.empty()
        progress_placeholder.empty()
        msg_placeholder.empty()

        # Set flag to show success screen on next render
        st.session_state["ptw_just_submitted"] = True
        st.session_state["ptw_is_submitting"] = False
        
        # Full visual reset (widgets) after SUCCESS only
        _reset_ptw_form_state()
        
        # Clear cache to ensure View Work Order shows updated status immediately
        st.cache_data.clear()
        st.session_state["s1_wo_last_df"] = None
        st.session_state["s1_wo_last_kpis"] = None
        
        # Rerun to show success screen
        st.rerun()

    except Exception as e:
        progress_placeholder.empty()
        status_placeholder.empty()
        msg_placeholder.empty()
        st.error(f"Error submitting PTW: {str(e)}")
        # Avoid showing stack traces to end users in production UI
        st.session_state["s1_ptw_last_file"] = None
        st.session_state["s1_ptw_last_permit_no"] = None
        st.session_state["ptw_is_submitting"] = False


def _extract_work_order_ids_from_ptw_row(r: dict) -> list[str]:
    """Multi-WO safe: prefer form_data['work_order_ids'], else legacy form_data['work_order_id'], else split permit_no."""
    if not isinstance(r, dict):
        return []
    fd = r.get("form_data") if isinstance(r.get("form_data"), dict) else {}
    if isinstance(fd, dict):
        ids = fd.get("work_order_ids")
        if isinstance(ids, list) and ids:
            out = [str(x).strip() for x in ids if str(x).strip()]
            if out:
                return out
        legacy = str(fd.get("work_order_id") or "").strip()
        if legacy:
            return [legacy]
    pn = str(r.get("permit_no") or "").strip()
    if pn and "-" in pn:
        parts = [p.strip() for p in pn.split("-") if p.strip()]
        return parts or [pn]
    return [pn] if pn else []


def _maybe_inject_closure_fields(*, form_data: dict, closure_dt: Any, receiver_name: str, holder_name: str, issuer_name: str) -> dict:
    """
    Ensure closure overlay fields exist in form_data.
    Uses the existing overlay tag sources (closure_* keys).
    """
    fd = dict(form_data or {})
    closure_dt_str = ""
    try:
        if closure_dt is not None:
            closure_dt_str = str(closure_dt)
    except Exception:
        closure_dt_str = ""

    # Only set if missing (do not overwrite)
    def _set_if_missing(k: str, v: str) -> None:
        if str(fd.get(k) or "").strip() == "" and str(v or "").strip() != "":
            fd[k] = v

    _set_if_missing("closure_receiver", receiver_name)
    _set_if_missing("closure_receiver_signature", receiver_name)
    _set_if_missing("closure_receiver_datetime", closure_dt_str)

    _set_if_missing("closure_holder", holder_name)
    _set_if_missing("closure_holder_signature", holder_name)
    _set_if_missing("closure_holder_datetime", closure_dt_str)

    _set_if_missing("closure_issuer", issuer_name)
    _set_if_missing("closure_issuer_signature", issuer_name)
    _set_if_missing("closure_issuer_datetime", closure_dt_str)

    return fd


@st.cache_data(show_spinner=False, ttl=60)
def _fetch_permit_closure_candidates(*, start_date: date, end_date: date) -> pd.DataFrame:
    """
    Return permits that are APPROVED or CLOSED, filtered by lifecycle clock: date_s3_approved in range.
    Multi-WO safe: aggregates PTW rows by ptw_requests.form_data['work_order_ids'].
    """
    sb = get_supabase_client(prefer_service_role=True)

    start_dt = datetime.combine(start_date, datetime.min.time())
    end_dt_excl = datetime.combine(end_date + timedelta(days=1), datetime.min.time())

    # Gate WOs by approved date range
    wo_gate = (
        sb.table(TABLE_WORK_ORDERS)
        .select(
            "work_order_id,site_name,location,equipment,"
            "date_s1_created,date_s2_forwarded,date_s3_approved,date_s2_rejected,date_s3_rejected,date_s1_closed"
        )
        .not_.is_("date_s3_approved", "null")
        .gte("date_s3_approved", start_dt.isoformat(sep=" ", timespec="seconds"))
        .lt("date_s3_approved", end_dt_excl.isoformat(sep=" ", timespec="seconds"))
        .execute()
    )
    wo_err = getattr(wo_gate, "error", None)
    if wo_err:
        raise RuntimeError(f"Failed to fetch approved work orders: {wo_err}")
    wo_rows: list[dict[str, Any]] = getattr(wo_gate, "data", None) or []
    gate_ids = {str(w.get("work_order_id") or "").strip() for w in wo_rows if str(w.get("work_order_id") or "").strip()}
    if not gate_ids:
        return pd.DataFrame()

    wo_lookup = {str(w["work_order_id"]).strip(): w for w in wo_rows if isinstance(w, dict) and w.get("work_order_id")}

    # Pull PTW requests and keep those that reference any gated id
    ptw_resp = (
        sb.table(TABLE_PTW_REQUESTS)
        .select("ptw_id,permit_no,site_name,created_at,created_by,form_data")
        .order("created_at", desc=True)
        .limit(1000)
        .execute()
    )
    ptw_err = getattr(ptw_resp, "error", None)
    if ptw_err:
        raise RuntimeError(f"Failed to fetch PTW requests: {ptw_err}")
    ptw_rows: list[dict[str, Any]] = getattr(ptw_resp, "data", None) or []

    out: list[dict[str, Any]] = []
    for r in ptw_rows:
        if not isinstance(r, dict):
            continue
        ids = _extract_work_order_ids_from_ptw_row(r)
        if not ids:
            continue
        if not (set(ids) & gate_ids):
            continue

        wo_local = [wo_lookup.get(i) for i in ids if i in wo_lookup]
        wo_local = [w for w in wo_local if isinstance(w, dict)]
        if not wo_local:
            continue

        # Permit-level lifecycle: compute per-WO derived and aggregate
        derived = [derive_ptw_status(w) for w in wo_local]
        sset = {str(x).strip().upper() for x in derived if str(x).strip()}
        if "REJECTED" in sset:
            permit_status = "REJECTED"
        elif sset and sset.issubset({"CLOSED"}):
            permit_status = "CLOSED"
        elif sset and sset.issubset({"APPROVED"}):
            permit_status = "APPROVED"
        else:
            # Not eligible for closure list
            continue

        # Aggregate approved/closed timestamps
        def _has(v) -> bool:
            return v is not None and str(v).strip() != ""

        appr_vals = [w.get("date_s3_approved") for w in wo_local if _has(w.get("date_s3_approved"))]
        clos_vals = [w.get("date_s1_closed") for w in wo_local if _has(w.get("date_s1_closed"))]
        appr_min = None
        clos_min = None
        try:
            ts = pd.to_datetime(appr_vals, errors="coerce")
            ts = ts.dropna() if hasattr(ts, "dropna") else ts
            if len(ts) > 0:
                appr_min = ts.min()
        except Exception:
            appr_min = appr_vals[0] if appr_vals else None
        try:
            ts = pd.to_datetime(clos_vals, errors="coerce")
            ts = ts.dropna() if hasattr(ts, "dropna") else ts
            if len(ts) > 0:
                clos_min = ts.min()
        except Exception:
            clos_min = clos_vals[0] if clos_vals else None

        fd = r.get("form_data") if isinstance(r.get("form_data"), dict) else {}
        receiver_name = str(fd.get("receiver_name") or fd.get("permit_receiver_name") or r.get("created_by") or "").strip()
        holder_name = str(fd.get("holder_name") or "").strip()
        issuer_name = str(fd.get("permit_issuer_name") or fd.get("issuer_name") or "").strip()

        out.append(
            {
                "ptw_id": r.get("ptw_id"),
                "permit_no": r.get("permit_no"),
                "site_name": r.get("site_name") or (wo_local[0].get("site_name") if wo_local else ""),
                "work_order_ids": ids,
                "approved_on": appr_min,
                "closed_on": clos_min,
                "status": permit_status,
                "receiver_name": receiver_name,
                "holder_name": holder_name,
                "issuer_name": issuer_name,
                "form_data": fd,
            }
        )

    df = pd.DataFrame(out)
    if df.empty:
        return df
    try:
        df = df.sort_values(by=["approved_on"], ascending=[False], na_position="last")
    except Exception:
        pass
    return df


def _render_permit_closure() -> None:
    """Render Permit Closure (S1 final stage) - updates work_orders.date_s1_closed and stores closure block in ptw_requests.form_data."""
    st.markdown("## Permit Closure")
    st.caption("Close S3-approved PTWs (final S1 stage). Closure is date-driven via `work_orders.date_s1_closed`.")

    col1, col2, col3 = st.columns([1.4, 1.4, 1.0], vertical_alignment="bottom")
    with col1:
        from_date = st.date_input("From Date", value=date.today() - timedelta(days=30), key="s1_close_from")
    with col2:
        to_date = st.date_input("To Date", value=date.today(), key="s1_close_to")
    with col3:
        if "s1_close_run_fetch" not in st.session_state:
            st.session_state["s1_close_run_fetch"] = False

        def _on_fetch() -> None:
            st.session_state["s1_close_run_fetch"] = True

        st.button("Fetch PTWs", type="primary", key="s1_close_fetch", on_click=_on_fetch)

    if "s1_close_df" not in st.session_state:
        st.session_state["s1_close_df"] = None

    if st.session_state.get("s1_close_run_fetch"):
        st.session_state["s1_close_run_fetch"] = False
        prog_slot = st.empty()
        prog = prog_slot.progress(0, text="Loading approved PTWs...")
        try:
            _smooth_progress(prog, 0, 25, text="Fetching candidates...")
            df = _fetch_permit_closure_candidates(start_date=from_date, end_date=to_date)
            st.session_state["s1_close_df"] = df
            _smooth_progress(prog, 25, 100, text="Ready")
        except Exception as e:
            st.error(f"Failed to fetch closure candidates: {e}")
            st.session_state["s1_close_df"] = pd.DataFrame()
        finally:
            prog_slot.empty()

    df = st.session_state.get("s1_close_df")
    if df is None:
        st.info("Select a date range and click 'Fetch PTWs' to load S3-approved permits.")
        return
    if df.empty:
        st.info("No approved/closed PTWs found for the selected date range.")
        return

    df_display = df.copy()
    df_display["status"] = df_display["status"].astype("string").fillna("").map(
        lambda s: DB_STATUS_TO_UI.get(str(s).strip().upper(), str(s))
    )
    # Color-code closure statuses (Approved = orange, Closed = green)
    disp = df_display[["permit_no", "site_name", "approved_on", "closed_on", "status"]].copy()

    def _style_closure_status(v: object) -> str:
        s = str(v or "").strip().lower()
        if s == "closed":
            return "background-color: #dcfce7; color: #065f46; font-weight: 800;"
        if s == "approved":
            return "background-color: #ffedd5; color: #7c2d12; font-weight: 800;"
        return ""

    styled = disp.style.map(_style_closure_status, subset=["status"]).set_table_styles(
        [{"selector": "th", "props": [("font-weight", "800"), ("color", "#0f172a")]}]
    )
    st.dataframe(styled, width="stretch", hide_index=True)

    st.divider()

    # Only approved (not yet closed) are closable
    open_df = df[df["status"].astype("string").str.upper() == "APPROVED"].copy()
    if open_df.empty:
        st.success("No permits pending closure in this date range.")
        return

    options = []
    for _, r in open_df.iterrows():
        pn = str(r.get("permit_no") or "").strip()
        site = str(r.get("site_name") or "").strip()
        appr = str(r.get("approved_on") or "").strip()
        options.append(f"{pn} | {site} | Approved: {appr}")

    selected_label = st.selectbox(
        "Select Permit to Close",
        options=["(select permit)"] + options,
        index=0,
        key="s1_close_select",
    )
    if not selected_label or selected_label == "(select permit)":
        st.info("Select a permit to view closure details.")
        return

    sel_pn = str(selected_label.split("|")[0]).strip()

    # Resolve selected row
    sel_row = None
    for _, r in open_df.iterrows():
        if str(r.get("permit_no") or "").strip() == sel_pn:
            sel_row = r
            break
    if sel_row is None:
        st.error("Selection could not be resolved. Please refetch.")
        return

    receiver_name = str(sel_row.get("receiver_name") or "")
    holder_name = str(sel_row.get("holder_name") or "")
    issuer_name = str(sel_row.get("issuer_name") or "")

    st.text_input("Permit Receiver Name", value=receiver_name, key="s1_close_receiver", disabled=True)
    st.text_input("Permit Holder Name", value=holder_name, key="s1_close_holder", disabled=True)
    st.text_input("Permit Issuer Name", value=issuer_name, key="s1_close_issuer", disabled=True)

    declared = st.checkbox(
        "I declare that the work area is restored to original state and all isolations are normalized.",
        key="s1_close_declaration",
    )

    if st.button("Close Permit", type="primary", key="s1_close_btn", disabled=not declared):
        with st.spinner("Closing permit..."):
            sb = get_supabase_client(prefer_service_role=True)
            now = datetime.now().isoformat(sep=" ", timespec="seconds")

            ptw_id = str(sel_row.get("ptw_id") or "").strip()
            wo_ids = sel_row.get("work_order_ids") or []
            wo_ids = [str(x).strip() for x in (wo_ids if isinstance(wo_ids, list) else []) if str(x).strip()]

            # Update all covered work orders (do not overwrite existing closures)
            for woid in wo_ids:
                sb.table(TABLE_WORK_ORDERS).update({"date_s1_closed": now}).eq("work_order_id", woid).is_(
                    "date_s1_closed", "null"
                ).execute()

            # Persist closure overlay block into ptw_requests.form_data for future downloads
            if ptw_id:
                get_resp = (
                    sb.table(TABLE_PTW_REQUESTS)
                    .select("form_data")
                    .eq("ptw_id", ptw_id)
                    .limit(1)
                    .execute()
                )
                cur = (getattr(get_resp, "data", None) or [{}])[0] or {}
                fd = cur.get("form_data") if isinstance(cur.get("form_data"), dict) else {}
                fd = _maybe_inject_closure_fields(
                    form_data=fd,
                    closure_dt=now,
                    receiver_name=receiver_name,
                    holder_name=holder_name,
                    issuer_name=issuer_name,
                )
                sb.table(TABLE_PTW_REQUESTS).update({"form_data": fd}).eq("ptw_id", ptw_id).execute()

            st.success("Permit closed successfully.")
            # Refresh cached lists
            st.cache_data.clear()
            st.session_state["s1_close_run_fetch"] = True
            st.rerun()


def _render_view_applied_ptw() -> None:
    """Render the View Applied PTW tab."""
    st.markdown("## View Applied PTW")
    st.caption("Browse and download previously submitted Permit To Work requests")

    # Filters
    col1, col2, col3 = st.columns([1.5, 1.5, 1])
    with col1:
        start_date = st.date_input(
            "From Date",
            value=date.today(),
            key="ptw_view_start"
        )
    with col2:
        end_date = st.date_input(
            "To Date",
            value=date.today(),
            key="ptw_view_end"
        )
    with col3:
        st.write("")  # Spacer
        st.write("")  # Spacer
        if "s1_view_ptw_run_fetch" not in st.session_state:
            st.session_state["s1_view_ptw_run_fetch"] = False
        def _on_fetch_ptws() -> None:
            st.session_state["s1_view_ptw_run_fetch"] = True
        fetch_btn = st.button("Fetch PTWs", type="primary", on_click=_on_fetch_ptws)

    # =========================================================================
    # FIX: Process fetch FIRST before any other rendering to prevent tab switch
    # The flag is checked and cleared immediately to prevent double-processing
    # =========================================================================
    should_fetch = st.session_state.get("s1_view_ptw_run_fetch", False)
    if should_fetch:
        # Clear flag FIRST to prevent re-triggering on subsequent reruns
        st.session_state["s1_view_ptw_run_fetch"] = False
        
        prog_slot = st.empty()
        prog = prog_slot.progress(0, text="Fetching PTWs...")
        try:
            _smooth_progress(prog, 0, 25, text="Validating date range...")
            _smooth_progress(prog, 25, 70, text="Fetching PTW requests...")
            df = fetch_ptw_requests(start_date=start_date, end_date=end_date)
            st.session_state["s1_ptw_view_df"] = df
            _smooth_progress(prog, 70, 100, text="Results ready")
        except Exception as e:
            st.error(f"Failed to fetch PTW requests: {e}")
            return
        finally:
            prog_slot.empty()
    
    # Show results if data exists (from current or previous fetch)
    if st.session_state.get("s1_ptw_view_df") is not None:

        df = st.session_state.get("s1_ptw_view_df")

        if df is None or df.empty:
            st.info("No PTW requests found for the selected date range.")
            return

        # Display summary
        st.metric("Total PTW Requests", len(df))

        # Format for display (lifecycle clock = work_orders.date_s1_created)
        cols = ["permit_no", "site_name", "status", "date_s1_created", "created_by"]
        for c in cols:
            if c not in df.columns:
                df[c] = pd.NA
        display_df = df[cols].copy()
        display_df["date_s1_created"] = pd.to_datetime(display_df["date_s1_created"], errors="coerce").dt.strftime("%Y-%m-%d %H:%M")
        display_df["status"] = display_df["status"].astype("string").fillna("").map(
            lambda s: DB_STATUS_TO_UI.get(str(s).strip().upper(), str(s))
        )
        display_df.columns = ["Permit No", "Site Name", "Status", "Submitted At", "Created By"]

        # Apply color-coded status styling
        styled = (
            display_df.style
            .map(_highlight_status, subset=["Status"])
            .set_table_styles([
                {"selector": "th", "props": [("font-weight", "800"), ("color", "#0f172a")]},
            ])
        )
        st.dataframe(styled, width="stretch", hide_index=True)

        # Select PTW to download (fragmented to avoid "whole page reload" feel)
        st.divider()
        st.subheader("Download PTW")

        _frag = getattr(st, "fragment", None)

        def _download_block() -> None:
            df_local = st.session_state.get("s1_ptw_view_df")
            if df_local is None or getattr(df_local, "empty", True):
                st.info("No PTW requests available.")
                return

            permit_options = df_local["permit_no"].tolist()
            selected_permit = st.selectbox(
                "Select Permit No to Download",
                options=permit_options,
                key="s1_view_selected_permit",
            )

            if "s1_view_ptw_pdf_bytes" not in st.session_state:
                st.session_state["s1_view_ptw_pdf_bytes"] = None
                st.session_state["s1_view_ptw_pdf_permit"] = None
                st.session_state["s1_view_ptw_pdf_approved_on"] = None

            def _on_generate_applied() -> None:
                st.session_state["s1_view_generate_clicked"] = True

            gen_clicked = st.button("Generate & Download", on_click=_on_generate_applied)

            if gen_clicked or st.session_state.get("s1_view_generate_clicked"):
                st.session_state["s1_view_generate_clicked"] = False
                try:
                    prog_slot = st.empty()
                    prog = prog_slot.progress(0, text="Generating report...")

                    # Get form data for selected permit
                    row = df_local[df_local["permit_no"] == selected_permit].iloc[0]
                    form_data = row.get("form_data", {})

                    if not form_data:
                        st.error("No form data found for this permit.")
                        return

                    # Inject approval timestamps (holder_datetime / issuer_datetime) from work_orders
                    # and apply APPROVED stamp if date_s3_approved is present.
                    _smooth_progress(prog, 0, 10, text="Checking approval status...")
                    wo_ids = row.get("work_order_ids")
                    if not isinstance(wo_ids, list) or not wo_ids:
                        wo_ids = _extract_work_order_ids_from_ptw_row({"permit_no": selected_permit, "form_data": form_data})
                    wo_ids = [str(x).strip() for x in (wo_ids or []) if str(x).strip()]
                    primary_wo = wo_ids[0] if wo_ids else selected_permit

                    updated_form, approval_times = _s1_inject_approval_times(form_data, primary_wo)
                    approved_on = approval_times.get("issuer_datetime") or ""

                    # If closed, inject closure overlay fields (uses stored names; does not overwrite)
                    closed_dt = None
                    try:
                        sb = get_supabase_client(prefer_service_role=True)
                        clos_resp = (
                            sb.table(TABLE_WORK_ORDERS)
                            .select("date_s1_closed")
                            .in_("work_order_id", wo_ids or [primary_wo])
                            .execute()
                        )
                        clos_rows = getattr(clos_resp, "data", None) or []
                        clos_vals = [r.get("date_s1_closed") for r in clos_rows if r.get("date_s1_closed") is not None]
                        if clos_vals:
                            closed_dt = pd.to_datetime(clos_vals, errors="coerce").min()
                    except Exception:
                        closed_dt = None

                    if closed_dt is not None and str(closed_dt).strip() != "" and str(closed_dt).strip().lower() != "nat":
                        recv = str(updated_form.get("receiver_name") or updated_form.get("permit_receiver_name") or "").strip()
                        holder = str(updated_form.get("holder_name") or "").strip()
                        issuer = str(updated_form.get("permit_issuer_name") or updated_form.get("issuer_name") or "").strip()
                        updated_form = _maybe_inject_closure_fields(
                            form_data=updated_form,
                            closure_dt=closed_dt,
                            receiver_name=recv,
                            holder_name=holder,
                            issuer_name=issuer,
                        )

                    # Generate document with attachments (uses S2's generator)
                    _smooth_progress(prog, 0, 85, text="Generating PDF with attachments...")
                    
                    def _progress_cb(pct: int, msg: str) -> None:
                        try:
                            prog.progress(int(85 * pct // 100), text=msg)
                        except Exception:
                            pass
                    
                    try:
                        from S2 import generate_ptw_pdf_with_attachments
                        # Use first work_order_id from permit's coverage
                        wo_ids = _extract_work_order_ids_from_ptw_row({"form_data": updated_form})
                        primary_wo_id = wo_ids[0] if wo_ids else selected_permit
                        pdf_bytes = generate_ptw_pdf_with_attachments(
                            form_data=updated_form,
                            work_order_id=primary_wo_id,
                            progress_callback=_progress_cb,
                        )
                    except Exception:
                        # Fallback: basic PDF without attachments
                        _smooth_progress(prog, 0, 40, text="Downloading template...")
                        template_bytes = _download_template_from_supabase()
                        _smooth_progress(prog, 40, 85, text="Generating PDF...")
                        doc_data = build_doc_data(updated_form)
                        pdf_bytes = generate_ptw_pdf(template_bytes, doc_data, progress_callback=_progress_cb)
                        pdf_bytes = _add_overlay_elements(pdf_bytes, updated_form)
                    
                    _smooth_progress(prog, 85, 100, text="Download ready")
                    prog_slot.empty()

                    st.session_state["s1_view_ptw_pdf_bytes"] = pdf_bytes
                    st.session_state["s1_view_ptw_pdf_permit"] = selected_permit
                    st.session_state["s1_view_ptw_pdf_approved_on"] = approved_on
                except Exception as e:
                    st.error(f"Failed to generate document: {e}")

            cached_pdf = st.session_state.get("s1_view_ptw_pdf_bytes")
            cached_permit = st.session_state.get("s1_view_ptw_pdf_permit")
            cached_approved_on = st.session_state.get("s1_view_ptw_pdf_approved_on") or ""
            if cached_pdf and cached_permit == selected_permit:
                st.download_button(
                    label=f"Download {selected_permit}.pdf",
                    data=cached_pdf,
                    file_name=f"{selected_permit}.pdf",
                    mime="application/pdf",
                    type="primary",
                    key=f"s1_view_cached_download_{selected_permit}",
                )

        if callable(_frag):
            _download_block_frag = _frag(_download_block)
            _download_block_frag()
        else:
            _download_block()
